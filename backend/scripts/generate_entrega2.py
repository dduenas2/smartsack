"""
Generador del documento de Entrega 2 del trabajo de grado SmartSack.

Produce `SmartSack_Entrega2_TrabajoDeGrado.docx` con 13 secciones + anexos,
reutilizando el estilo del Anteproyecto y la Entrega 1.

El script corre dentro del contenedor backend, donde `/app` está bind-mounted a
`./backend/` en el host. El archivo se escribe en `/docs/` dentro del
contenedor (NO está bind-mounted) y debe extraerse al host con `docker cp`:

    docker compose exec backend python -m scripts.generate_entrega2
    docker cp smartsack_backend:/docs/SmartSack_Entrega2_TrabajoDeGrado.docx \\
        docs/SmartSack_Entrega2_TrabajoDeGrado.docx
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


OUT_PATH = Path("/app/../docs/SmartSack_Entrega2_TrabajoDeGrado.docx")


# -----------------------------------------------------------------------------
# Helpers de estilo
# -----------------------------------------------------------------------------
def set_cell_shading(cell, hex_color: str) -> None:
    """Pinta la celda con un color de fondo (sombreado)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_table(doc, headers, rows, *, header_color: str = "1F4E79"):
    """Inserta una tabla con cabecera coloreada y autofit."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.autofit = True

    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        set_cell_shading(hdr[i], header_color)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for i, row in enumerate(rows, start=1):
        cells = t.rows[i].cells
        for j, val in enumerate(row):
            cells[j].text = ""
            p = cells[j].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return t


def add_para(doc, text: str, *, bold: bool = False, italic: bool = False, size: int = 11, justify: bool = True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def add_bullets(doc, items, style: str = "List Bullet"):
    for item in items:
        p = doc.add_paragraph(style=style)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(item).font.size = Pt(11)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(item).font.size = Pt(11)


def add_heading(doc, text: str, level: int):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    return h


def add_pagebreak(doc):
    doc.add_page_break()


# -----------------------------------------------------------------------------
# PORTADA
# -----------------------------------------------------------------------------
def write_cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("INSTITUCIÓN UNIVERSITARIA POLITÉCNICO GRANCOLOMBIANO").bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("FACULTAD DE INGENIERÍA, DISEÑO E INNOVACIÓN — FIDI").bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("PROGRAMA DE INGENIERÍA DE SOFTWARE").bold = True

    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TRABAJO DE GRADO — ENTREGA 2")
    r.bold = True
    r.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Análisis, diseño y avance de implementación")
    r.italic = True
    r.font.size = Pt(13)

    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "SmartSack: Plataforma inteligente de gestión de producción con Digital Twin, "
        "analítica predictiva y asistente conversacional con IA para plantas de "
        "fabricación de sacos de papel"
    )
    r.bold = True
    r.font.size = Pt(13)

    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("David Alejandro Dueñas Castrillon").bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Código: 100320124")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Asesor: Carlos David Seligmann Trujillo")

    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Medellín, Colombia")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(date.today().strftime("%d de %B de %Y").lower().replace("january", "enero")
              .replace("february", "febrero").replace("march", "marzo").replace("april", "abril")
              .replace("may", "mayo").replace("june", "junio").replace("july", "julio")
              .replace("august", "agosto").replace("september", "septiembre").replace("october", "octubre")
              .replace("november", "noviembre").replace("december", "diciembre"))
    add_pagebreak(doc)


# -----------------------------------------------------------------------------
# RESUMEN EJECUTIVO
# -----------------------------------------------------------------------------
def write_resumen(doc):
    add_heading(doc, "RESUMEN DE LA ENTREGA 2", level=1)
    add_para(
        doc,
        "Este documento corresponde a la segunda de tres entregas oficiales del trabajo de "
        "grado SmartSack. La Entrega 1 fundamentó el proyecto desde la perspectiva del "
        "problema, los objetivos, el marco teórico, el estado del arte y la metodología, "
        "mientras que la Entrega 3 documentará la validación con usuarios reales, los "
        "resultados experimentales y las conclusiones. La presente Entrega 2 cubre la fase "
        "intermedia y se concentra en transformar la propuesta en una solución diseñada y "
        "parcialmente construida."
    )
    add_para(
        doc,
        "Específicamente, en este documento se presentan: (1) la especificación de "
        "requisitos funcionales y no funcionales derivados de las brechas operativas "
        "identificadas; (2) el modelo de casos de uso con sus actores y flujos; (3) la "
        "arquitectura del sistema con diagramas de componentes y de despliegue; (4) el "
        "modelo de datos relacional implementado en PostgreSQL; (5) el diseño de la "
        "interfaz con descripción de cada vista; (6) el diseño del modelo de Machine "
        "Learning siguiendo el ciclo CRISP-DM; (7) el diseño del proceso ETL y la captura "
        "de eventos en tiempo real; (8) el diseño del asistente conversacional con la API "
        "de Claude y LangChain; (9) la estrategia de pruebas con su nivel actual de "
        "cobertura; (10) el avance verificable de la implementación organizado por "
        "objetivo específico; y (11) los riesgos identificados con sus mitigaciones."
    )
    add_para(
        doc,
        "El proyecto, en su estado actual, cuenta con una plataforma operativa contenedorizada "
        "con Docker Compose que reúne cinco servicios (PostgreSQL, Redis, backend FastAPI, "
        "frontend React y reverse proxy Nginx). Se han implementado catorce entidades de "
        "negocio en el modelo de datos, cuarenta y seis endpoints REST documentados con "
        "Swagger/OpenAPI, ocho vistas funcionales en el frontend, un canal WebSocket para "
        "el Digital Twin del supervisor, un proceso ETL multi-formato, un modelo XGBoost "
        "entrenado y serializado con treinta variables, ocho herramientas de function "
        "calling para el chatbot, y un total de ciento noventa y siete pruebas automatizadas "
        "(ciento sesenta y siete en backend y treinta en frontend) que ejecutan en cada "
        "ciclo de integración. Sobre esta base, la Entrega 3 ejecutará la validación "
        "experimental definida en el diseño cuasi-experimental pre-post de la Entrega 1."
    )
    add_para(
        doc,
        "Palabras clave: análisis de requisitos, casos de uso, arquitectura de software, "
        "modelo entidad-relación, CRISP-DM, function calling, pruebas automatizadas, "
        "Docker Compose, FastAPI, React, XGBoost, LangChain, Digital Twin.",
        italic=True,
    )
    add_pagebreak(doc)


# -----------------------------------------------------------------------------
# 1. INTRODUCCIÓN Y CONTEXTO
# -----------------------------------------------------------------------------
def write_intro(doc):
    add_heading(doc, "1. INTRODUCCIÓN Y CONTEXTO", level=1)

    add_heading(doc, "1.1. Antecedentes del proyecto", level=2)
    add_para(
        doc,
        "El proyecto SmartSack fue formulado en el Anteproyecto entregado en marzo de 2026 "
        "como una plataforma web inteligente complementaria a los sistemas ERP existentes "
        "en plantas de fabricación de sacos de papel. La Entrega 1, fechada en abril de "
        "2026, profundizó el planteamiento del problema identificando cuatro brechas "
        "operativas críticas: ausencia de visibilidad en tiempo real, comunicación "
        "fragmentada entre áreas funcionales, falta de trazabilidad integral y ausencia "
        "de capacidad predictiva sobre los datos históricos disponibles en el ERP."
    )
    add_para(
        doc,
        "Frente a estas brechas, la propuesta integra tres pilares tecnológicos que "
        "operan de manera complementaria: un Digital Twin de la línea de producción "
        "con vistas diferenciadas por rol; un motor de analítica predictiva basado en "
        "Machine Learning para la predicción de retrasos en órdenes; y un asistente "
        "conversacional impulsado por un modelo de lenguaje de gran escala que permite "
        "consultas en lenguaje natural mediante function calling. Estos tres pilares se "
        "construyen sobre una arquitectura unificada en Python (FastAPI) con frontend en "
        "React, base de datos PostgreSQL, caché Redis y reverse proxy Nginx, todo "
        "contenedorizado con Docker Compose."
    )

    add_heading(doc, "1.2. Propósito de esta entrega", level=2)
    add_para(
        doc,
        "La Entrega 2 cumple cuatro propósitos académicos articulados. Primero, traduce "
        "el planteamiento conceptual de la Entrega 1 a una especificación técnica "
        "trazable: cada brecha del problema se transforma en uno o más requisitos "
        "funcionales, cada requisito funcional se vincula a uno o más casos de uso, y "
        "cada caso de uso se materializa en endpoints REST y componentes de interfaz "
        "concretos. Segundo, documenta las decisiones de diseño que dan forma a la "
        "implementación: la elección de la arquitectura cliente-servidor contenedorizada, "
        "el modelo entidad-relación, los algoritmos de Machine Learning evaluados, las "
        "tecnologías de tiempo real (WebSockets) y los mecanismos de seguridad (JWT con "
        "roles). Tercero, presenta el avance verificable del prototipo construido durante "
        "los sprints de las semanas 4 a 8, con métricas concretas de cobertura de pruebas, "
        "número de endpoints funcionales y módulos integrados. Cuarto, identifica los "
        "riesgos remanentes y las mitigaciones planificadas para la Entrega 3."
    )

    add_heading(doc, "1.3. Alcance de esta entrega", level=2)
    add_para(
        doc,
        "Quedan dentro del alcance de la Entrega 2 los siguientes artefactos: "
        "especificación de requisitos funcionales y no funcionales; modelo de casos de "
        "uso con descripción de actores y flujos principales; diagrama de arquitectura "
        "del sistema y diagrama de despliegue; modelo entidad-relación con diccionario de "
        "datos; descripción de la interfaz de usuario por vista; diseño del pipeline de "
        "Machine Learning siguiendo CRISP-DM, con métricas obtenidas sobre datos "
        "sintéticos; diseño del proceso ETL y de la captura de eventos en tiempo real; "
        "diseño del asistente conversacional con sus ocho herramientas de function "
        "calling; estrategia de pruebas con cobertura actual; reporte de avance por "
        "objetivo específico; e identificación de riesgos."
    )
    add_para(
        doc,
        "Quedan fuera del alcance de la Entrega 2 y se reservan explícitamente para la "
        "Entrega 3: la aplicación del cuestionario SUS y la encuesta de satisfacción a "
        "usuarios reales; el registro cronometrado pre-post de tiempos de respuesta; el "
        "reentrenamiento del modelo de Machine Learning con datos reales del ERP "
        "(actualmente se trabaja con datos sintéticos generados por el script de seed); "
        "la conexión efectiva del chatbot con la API de Claude (que requiere clave "
        "comercial activa); el análisis estadístico inferencial pre-post; los manuales "
        "técnico, de despliegue y de usuario; y la preparación de la sustentación. "
        "Esta separación garantiza que cada entrega tenga un enfoque diferenciado y "
        "evaluable de manera independiente."
    )

    add_heading(doc, "1.4. Estructura del documento", level=2)
    add_para(
        doc,
        "El documento se organiza en trece secciones principales y cuatro anexos. La "
        "sección 2 presenta la especificación de requisitos. La sección 3 describe el "
        "modelo de casos de uso. La sección 4 expone la arquitectura del sistema. La "
        "sección 5 detalla el modelo de datos. La sección 6 documenta el diseño de la "
        "interfaz de usuario. La sección 7 cubre el diseño del modelo de Machine "
        "Learning. La sección 8 describe el diseño del ETL y la captura de eventos en "
        "tiempo real. La sección 9 explica el diseño del asistente conversacional. La "
        "sección 10 presenta la estrategia de pruebas. La sección 11 reporta el avance "
        "de implementación por objetivo específico. La sección 12 identifica riesgos y "
        "mitigaciones. La sección 13 contiene la bibliografía. Los anexos consolidan el "
        "diccionario de datos, las fichas detalladas de los casos de uso principales, "
        "los snippets de código relevantes y el inventario completo de endpoints REST."
    )
    add_pagebreak(doc)


# -----------------------------------------------------------------------------
# 2. ANÁLISIS DE REQUISITOS
# -----------------------------------------------------------------------------
def write_requisitos(doc):
    add_heading(doc, "2. ANÁLISIS DE REQUISITOS", level=1)
    add_para(
        doc,
        "Esta sección consolida los requisitos funcionales y no funcionales del sistema "
        "SmartSack. Cada requisito se identifica con un código único, se justifica desde "
        "la brecha operativa que pretende cerrar (sección 2.2 de la Entrega 1) y se "
        "vincula con los casos de uso correspondientes presentados en la sección 3. La "
        "trazabilidad se mantiene en una matriz transversal que se incluye al final de "
        "esta sección."
    )

    add_heading(doc, "2.1. Stakeholders y actores del sistema", level=2)
    add_para(
        doc,
        "El sistema reconoce cuatro tipos de actores con diferentes niveles de privilegio "
        "y casos de uso característicos. La diferenciación se implementa a través de "
        "JSON Web Tokens (JWT) firmados con HMAC-SHA256, que codifican el rol y, en el "
        "caso del operario, la máquina asignada."
    )
    add_table(
        doc,
        ["Actor", "Descripción", "Casos de uso típicos"],
        [
            ["Operario", "Trabajador asignado a una máquina específica que registra eventos en tiempo real (paradas, cambios de formato, incidencias, fin de operación) desde el PC de su máquina.", "CU-04 Iniciar operación, CU-05 Registrar evento, CU-06 Reportar avance, CU-07 Cerrar operación"],
            ["Supervisor", "Responsable de una línea o turno que monitorea el estado de la planta, gestiona órdenes y consulta predicciones de retraso.", "CU-08 Visualizar Digital Twin, CU-09 Consultar dashboard, CU-10 Consultar predicción, CU-11 Conversar con asistente"],
            ["Administrador", "Encargado de la configuración del sistema, gestión de usuarios y máquinas, y carga de datos por ETL.", "CU-12 Cargar archivo ETL, CU-13 Gestionar usuarios, CU-14 Gestionar máquinas, CU-15 Auditar acciones"],
            ["Sistema", "Procesos automáticos: cálculo periódico de OEE, refresco del modelo ML, ping de WebSocket, generación de predicciones para órdenes activas.", "Tareas internas no expuestas al usuario."],
        ],
    )

    add_heading(doc, "2.2. Requisitos funcionales (RF)", level=2)
    add_para(
        doc,
        "Los siguientes requisitos describen las capacidades observables que el sistema "
        "debe ofrecer. Cada requisito está vinculado a la brecha que resuelve y al "
        "objetivo específico de la Entrega 1 al que contribuye."
    )
    rf_rows = [
        ["RF-01", "Autenticación", "El sistema debe autenticar usuarios mediante usuario/contraseña y emitir un JWT firmado que codifique el rol y, si aplica, la máquina asignada.", "Brecha 2", "Obj. 1, 6"],
        ["RF-02", "Autorización por roles", "El sistema debe rechazar (HTTP 403) cualquier acceso a un recurso para el cual el rol del JWT no esté autorizado.", "Brecha 2", "Obj. 1, 6"],
        ["RF-03", "Vista de operario", "El operario debe poder ver el estado de su máquina, la operación en curso, la cola de operaciones siguientes y el historial reciente de eventos en tiempo real.", "Brecha 1", "Obj. 3"],
        ["RF-04", "Registro rápido de eventos", "El operario debe poder registrar paradas, cambios de formato, incidencias y reportes parciales de avance con un máximo de tres clics.", "Brecha 1, 3", "Obj. 3"],
        ["RF-05", "Cerrar operación", "El operario debe poder cerrar la operación activa indicando cantidad producida y kg de scrap con su razón.", "Brecha 1, 3", "Obj. 3"],
        ["RF-06", "Vista de supervisor (Digital Twin)", "El supervisor debe ver un mapa de planta con todas las máquinas codificadas por color según su estado, actualizado por WebSocket en menos de un segundo desde el cambio de estado.", "Brecha 1", "Obj. 3"],
        ["RF-07", "Dashboard de KPIs", "El sistema debe ofrecer un dashboard con OEE de planta, OEE de ayer, descomposición A/P/Q, producción del día, ranking de máquinas, alertas, scrap Pareto, yield por operación, y WIP en línea.", "Brecha 1, 4", "Obj. 4"],
        ["RF-08", "Predicción de retrasos", "El sistema debe predecir la probabilidad de retraso de cada orden activa con un modelo XGBoost entrenado sobre el histórico, y exponer las predicciones por endpoint REST.", "Brecha 4", "Obj. 4"],
        ["RF-09", "Asistente conversacional", "El sistema debe responder consultas en lenguaje natural sobre datos de producción, OEE, paradas, scrap, yield y WIP, redirigiendo al modelo Claude cuando hay credenciales o a un router heurístico de fallback en caso contrario.", "Brecha 2", "Obj. 5"],
        ["RF-10", "ETL desde CSV", "El administrador debe poder subir archivos CSV de cuatro tipos (production_orders, confirmations, materials, shipments) y el sistema debe validar columnas, detectar duplicados, registrar errores y producir un informe de carga.", "Brecha 4", "Obj. 2"],
        ["RF-11", "Gestión de usuarios y máquinas", "El administrador debe poder crear, listar, editar y desactivar usuarios y máquinas, así como asignar operarios a máquinas y resetear contraseñas.", "—", "Obj. 1, 6"],
        ["RF-12", "Auditoría de acciones administrativas", "El sistema debe registrar en una tabla de auditoría toda acción administrativa de creación, modificación o eliminación, con autor, marca de tiempo y diff antes/después.", "Brecha 3", "Obj. 1"],
        ["RF-13", "Trazabilidad de orden", "El sistema debe permitir consultar el recorrido completo de una orden, mostrando todas las operaciones, eventos, predicciones y registros de calidad asociados.", "Brecha 3", "Obj. 3"],
        ["RF-14", "Consultar plantillas CSV", "El sistema debe ofrecer la descarga de plantillas vacías (cabeceras) por cada tipo de archivo ETL aceptado.", "—", "Obj. 2"],
        ["RF-15", "Health check del sistema", "El sistema debe exponer un endpoint público de health (postgres, redis, anthropic, modelo ML) y un endpoint detallado para administradores.", "—", "Obj. 6"],
    ]
    add_table(
        doc,
        ["ID", "Nombre", "Descripción", "Brecha", "Objetivo"],
        rf_rows,
    )

    add_heading(doc, "2.3. Requisitos no funcionales (RNF)", level=2)
    add_para(
        doc,
        "Los requisitos no funcionales describen características de calidad transversales: "
        "cómo debe comportarse el sistema, no qué debe hacer. Se organizan en seis "
        "categorías habituales: rendimiento, seguridad, usabilidad, mantenibilidad, "
        "portabilidad y disponibilidad."
    )
    rnf_rows = [
        ["RNF-01", "Rendimiento", "El tiempo de respuesta del 95% de las consultas REST debe ser inferior a 500 ms, medido sobre 6 meses de datos sintéticos en entorno de desarrollo (PostgreSQL local, sin réplica).", "Latencia P95 ≤ 500 ms"],
        ["RNF-02", "Tiempo real", "Las actualizaciones del Digital Twin del supervisor deben propagarse a todos los clientes WebSocket conectados en menos de un segundo desde el cambio de estado en la base de datos.", "Latencia WS ≤ 1 s"],
        ["RNF-03", "Seguridad — autenticación", "Todas las contraseñas deben almacenarse cifradas con bcrypt (cost factor ≥ 12). Los JWT deben firmarse con HMAC-SHA256 y un secreto de al menos 64 caracteres aleatorios.", "Bcrypt, JWT HS256"],
        ["RNF-04", "Seguridad — transporte", "En despliegue de producción, todo el tráfico debe estar bajo HTTPS terminado en Nginx; los WebSockets deben usar wss://.", "TLS 1.2+"],
        ["RNF-05", "Seguridad — secretos", "Ningún valor sensible (claves, contraseñas, tokens) puede estar embebido en el código fuente. Toda configuración sensible se inyecta vía variables de entorno (.env).", "0 secretos en repo"],
        ["RNF-06", "Usabilidad", "La interfaz debe ser operable desde cualquier navegador moderno (Chrome 100+, Edge 100+, Firefox 100+) en el monitor estándar de los PCs de planta (1366×768) sin scroll horizontal.", "SUS ≥ 70 (E3)"],
        ["RNF-07", "Internacionalización", "Toda cadena visible para el usuario debe estar en español; los identificadores de código (variables, funciones, archivos) en inglés.", "Documentación de la convención"],
        ["RNF-08", "Mantenibilidad — pruebas", "El backend debe tener al menos 150 tests automatizados ejecutables con pytest. El frontend debe tener al menos 25 tests con Vitest.", "197 tests (logrado)"],
        ["RNF-09", "Mantenibilidad — documentación", "Cada archivo del backend y del frontend debe iniciar con un docstring que describa su propósito; cada función pública debe documentar parámetros y retorno cuando no sea evidente del nombre.", "100% archivos documentados"],
        ["RNF-10", "Portabilidad", "El sistema completo debe levantarse en cualquier máquina con Docker Compose v2 ejecutando `docker compose up -d --build` desde el directorio raíz, en menos de 5 minutos en hardware estándar (4 vCPU, 8 GB RAM).", "Tiempo de boot ≤ 5 min"],
        ["RNF-11", "Disponibilidad", "Los contenedores deben tener policy `restart: unless-stopped` y healthchecks definidos en docker-compose.yml para PostgreSQL, Redis y backend.", "Auto-restart configurado"],
        ["RNF-12", "Auditoría", "Cualquier acción de creación, modificación o eliminación realizada por un administrador debe quedar registrada en `admin_audit_log` con autor, timestamp e información antes/después.", "Tabla auditoría operacional"],
        ["RNF-13", "Compatibilidad ML", "El modelo entrenado debe poder cargarse en cualquier instancia con la misma versión de scikit-learn y XGBoost mediante el manifest JSON adjunto al artefacto joblib.", "Manifest versionado"],
    ]
    add_table(
        doc,
        ["ID", "Categoría", "Descripción", "Métrica/Verificación"],
        rnf_rows,
    )

    add_heading(doc, "2.4. Matriz de trazabilidad brecha → requisito → caso de uso", level=2)
    add_para(
        doc,
        "La matriz siguiente garantiza que cada brecha operativa identificada en la "
        "Entrega 1 da origen a al menos un requisito funcional, y que cada requisito "
        "funcional se cubre con al menos un caso de uso. Esta trazabilidad es esencial "
        "para la evaluación del jurado y para la fase de validación de la Entrega 3."
    )
    add_table(
        doc,
        ["Brecha (E1 §2.2)", "Requisitos cubiertos", "Casos de uso"],
        [
            ["B1 — Ausencia de visibilidad en tiempo real", "RF-03, RF-04, RF-05, RF-06, RF-07", "CU-04, CU-05, CU-06, CU-07, CU-08, CU-09"],
            ["B2 — Comunicación fragmentada", "RF-01, RF-02, RF-09", "CU-01, CU-11"],
            ["B3 — Falta de trazabilidad integral", "RF-04, RF-05, RF-12, RF-13", "CU-05, CU-15, todos los CU de operario"],
            ["B4 — Ausencia de capacidad predictiva", "RF-07, RF-08, RF-10", "CU-09, CU-10, CU-12"],
        ],
    )
    add_pagebreak(doc)


# -----------------------------------------------------------------------------
# 3. CASOS DE USO
# -----------------------------------------------------------------------------
def write_casos_uso(doc):
    add_heading(doc, "3. MODELO DE CASOS DE USO", level=1)
    add_para(
        doc,
        "Esta sección presenta el modelo de casos de uso del sistema SmartSack. Se "
        "definen quince casos de uso principales agrupados por actor; siete de ellos se "
        "documentan con ficha detallada (precondiciones, flujos básico y alternativos, "
        "postcondiciones) en el Anexo B. Los restantes se consideran casos de uso de "
        "soporte cuya descripción a alto nivel es suficiente para el alcance de esta "
        "entrega."
    )

    add_heading(doc, "3.1. Diagrama de casos de uso (descripción)", level=2)
    add_para(
        doc,
        "El diagrama de casos de uso del sistema se estructura en torno a cuatro "
        "actores conectados con quince casos de uso. Cada actor tiene una zona de "
        "responsabilidad clara, con relaciones <<include>> y <<extend>> para casos "
        "compartidos. La representación gráfica recomendada (a producir en draw.io o "
        "PlantUML) sigue la siguiente especificación textual:"
    )
    add_bullets(
        doc,
        [
            "Operario: conectado a CU-01 (Autenticarse), CU-04 (Iniciar operación), CU-05 (Registrar evento), CU-06 (Reportar avance) y CU-07 (Cerrar operación). CU-04, CU-05, CU-06 y CU-07 incluyen <<include>> CU-01.",
            "Supervisor: conectado a CU-01, CU-08 (Visualizar Digital Twin), CU-09 (Consultar dashboard), CU-10 (Consultar predicción) y CU-11 (Conversar con asistente). CU-09 incluye <<include>> CU-08.",
            "Administrador: conectado a CU-01, CU-12 (Cargar archivo ETL), CU-13 (Gestionar usuarios), CU-14 (Gestionar máquinas) y CU-15 (Auditar acciones). Hereda los CU del Supervisor.",
            "Sistema (actor automático): vinculado a CU-02 (Calcular OEE diario) y CU-03 (Refrescar predicciones de órdenes activas), invocados por scripts CLI o tareas programadas.",
        ],
    )
    add_para(
        doc,
        "El diagrama se coloca en el Anexo B como imagen recomendada; su construcción "
        "se delega a la herramienta gráfica preferida del estudiante (draw.io, "
        "Lucidchart, PlantUML o Visio). La especificación textual permite reproducir el "
        "diagrama de manera unívoca."
    )

    add_heading(doc, "3.2. Inventario de casos de uso", level=2)
    add_table(
        doc,
        ["ID", "Nombre", "Actor primario", "Tipo", "Detalle en"],
        [
            ["CU-01", "Autenticarse en el sistema", "Todos los usuarios", "Esencial", "Anexo B.1"],
            ["CU-02", "Calcular OEE diario", "Sistema", "Soporte", "Sección 3.3"],
            ["CU-03", "Refrescar predicciones de órdenes activas", "Sistema", "Soporte", "Sección 3.3"],
            ["CU-04", "Iniciar operación", "Operario", "Esencial", "Anexo B.2"],
            ["CU-05", "Registrar evento de máquina", "Operario", "Esencial", "Anexo B.3"],
            ["CU-06", "Reportar avance parcial", "Operario", "Esencial", "Anexo B.4"],
            ["CU-07", "Cerrar operación", "Operario", "Esencial", "Sección 3.3"],
            ["CU-08", "Visualizar Digital Twin", "Supervisor", "Esencial", "Anexo B.5"],
            ["CU-09", "Consultar dashboard de KPIs", "Supervisor", "Esencial", "Anexo B.6"],
            ["CU-10", "Consultar predicción de orden", "Supervisor", "Esencial", "Sección 3.3"],
            ["CU-11", "Conversar con asistente IA", "Supervisor", "Esencial", "Anexo B.7"],
            ["CU-12", "Cargar archivo ETL", "Administrador", "Esencial", "Sección 3.3"],
            ["CU-13", "Gestionar usuarios", "Administrador", "Soporte", "Sección 3.3"],
            ["CU-14", "Gestionar máquinas", "Administrador", "Soporte", "Sección 3.3"],
            ["CU-15", "Auditar acciones del sistema", "Administrador", "Soporte", "Sección 3.3"],
        ],
    )

    add_heading(doc, "3.3. Resúmenes de casos de uso de soporte", level=2)
    add_para(doc, "CU-02 — Calcular OEE diario", bold=True)
    add_para(
        doc,
        "Disparado por un script CLI o por el primer registro del día. El sistema "
        "agrupa eventos por máquina y turno, computa Disponibilidad como tiempo "
        "operativo / tiempo planificado, Rendimiento como cantidad real / cantidad "
        "ideal, y Calidad como buenos / total, y persiste el resultado en "
        "`oee_records` con la marca de fecha. Postcondición: existe un registro de "
        "OEE por (máquina, turno, fecha) para todos los turnos del día anterior."
    )
    add_para(doc, "CU-03 — Refrescar predicciones de órdenes activas", bold=True)
    add_para(
        doc,
        "Invocado vía endpoint POST /api/predictions/predict-active. Para cada orden con "
        "estado pending, in_progress o delayed, el sistema construye el vector de 30 "
        "features definidos en `ml/features.py`, aplica el modelo XGBoost cargado en "
        "memoria y persiste una nueva fila en `ml_predictions` con probabilidad y "
        "horas estimadas de retraso. Postcondición: cada orden activa tiene una "
        "predicción cuya antigüedad es menor a 24 horas."
    )
    add_para(doc, "CU-07 — Cerrar operación", bold=True)
    add_para(
        doc,
        "El operario, desde la vista OperatorView, marca la operación en curso como "
        "completada indicando cantidad producida y, opcionalmente, kg de scrap con su "
        "razón (quality_defect, setup_loss, material_break, other). El sistema "
        "actualiza `order_operations.status = COMPLETED`, registra `actual_end`, y "
        "evalúa si la siguiente operación de la orden puede pasar a READY. "
        "Postcondición: la operación está cerrada y su evento se difunde por WebSocket "
        "al supervisor."
    )
    add_para(doc, "CU-10 — Consultar predicción de orden", bold=True)
    add_para(
        doc,
        "El supervisor solicita la predicción para una orden específica vía POST "
        "/api/predictions/predict/{order_id}. El sistema responde con probabilidad de "
        "retraso, horas estimadas y la versión del modelo utilizado. La interfaz "
        "muestra la última predicción registrada y permite forzar un recálculo."
    )
    add_para(doc, "CU-12 — Cargar archivo ETL", bold=True)
    add_para(
        doc,
        "El administrador, desde la vista ETL, selecciona el tipo de archivo "
        "(production_orders, confirmations, materials, shipments) y arrastra un CSV. "
        "El sistema valida cabeceras, lee el archivo con Pandas, detecta duplicados, "
        "inserta o actualiza filas en PostgreSQL, registra el resultado en `etl_loads` "
        "(filas totales, insertadas, actualizadas, omitidas, fallidas, duración) y "
        "devuelve un informe que se muestra en la UI."
    )
    add_para(doc, "CU-13 — Gestionar usuarios", bold=True)
    add_para(
        doc,
        "Operaciones CRUD sobre la entidad User a través de los endpoints /api/users. "
        "Incluye creación, edición, asignación de máquina al operario, reset de "
        "contraseña y desactivación. Toda acción queda registrada en "
        "`admin_audit_log`."
    )
    add_para(doc, "CU-14 — Gestionar máquinas", bold=True)
    add_para(
        doc,
        "Operaciones CRUD sobre la entidad Machine. Permite añadir nuevas máquinas a la "
        "planta (por ejemplo, ampliaciones de línea), modificar su estado y, "
        "excepcionalmente, eliminarlas. Las altas y bajas modifican el conjunto de "
        "tiles del Digital Twin del supervisor en tiempo real."
    )
    add_para(doc, "CU-15 — Auditar acciones del sistema", bold=True)
    add_para(
        doc,
        "El administrador consulta /api/admin/audit con filtros por actor, acción, "
        "tipo de entidad o rango de fechas. La respuesta muestra los diff completos "
        "(JSON antes y después) que dejan trazabilidad legal de cualquier modificación "
        "realizada por un administrador."
    )
    add_pagebreak(doc)


# -----------------------------------------------------------------------------
# 4. ARQUITECTURA DEL SISTEMA
# -----------------------------------------------------------------------------
def write_arquitectura(doc):
    add_heading(doc, "4. ARQUITECTURA DEL SISTEMA", level=1)
    add_para(
        doc,
        "Esta sección expone la arquitectura del sistema desde tres vistas "
        "complementarias: la vista de componentes (qué módulos lo conforman y cómo se "
        "comunican), la vista de despliegue (en qué contenedores corren y cómo se "
        "orquestan) y la vista de seguridad (cómo se protegen los activos)."
    )

    add_heading(doc, "4.1. Vista de componentes", level=2)
    add_para(
        doc,
        "SmartSack adopta una arquitectura cliente-servidor de tres capas con un canal "
        "adicional de tiempo real. Las cinco unidades de despliegue son: el "
        "frontend SPA (React + Vite + Tailwind), el backend (FastAPI + SQLAlchemy + "
        "PyJWT + LangChain + scikit-learn/XGBoost), la base de datos relacional "
        "(PostgreSQL 16), la caché en memoria (Redis 7) y el reverse proxy (Nginx). "
        "El frontend consume al backend exclusivamente a través de Nginx; nunca le "
        "habla a la base de datos directamente. Los WebSockets viajan por la misma "
        "ruta del proxy bajo el path /ws/."
    )
    add_para(
        doc,
        "El backend se organiza internamente en cinco subpaquetes: routers (la capa "
        "expuesta vía HTTP/WS), services (la lógica de negocio), models (las entidades "
        "SQLAlchemy y enumeraciones), schemas (las DTO Pydantic) y los dominios "
        "transversales chat, websocket, etl, ml. Cada router invoca a su servicio "
        "correspondiente; los servicios usan SQLAlchemy para persistencia y emiten "
        "eventos al WebSocket manager cuando proceden."
    )
    add_para(
        doc,
        "El módulo ml/ contiene el pipeline de Machine Learning y se organiza en "
        "features.py (construcción del vector de 30 features), train.py (entrenamiento "
        "comparativo de Random Forest y XGBoost con validación cruzada estratificada de "
        "5 folds y GridSearchCV) y models/ (artefacto serializado con joblib + manifest "
        "JSON con métricas, hiperparámetros y feature importance)."
    )
    add_para(
        doc,
        "El módulo chat/tools.py es agnóstico al frontend de IA: expone ocho funciones "
        "Python con sus schemas Anthropic-tools-compatible. El servicio "
        "services/chat_service.py orquesta dos modos: (1) modo LLM, que delega en "
        "ChatAnthropic + LangChain con bind_tools para iterar hasta cuatro veces "
        "tool→result→tool, y (2) modo fallback, un router heurístico por palabras "
        "clave que selecciona la tool correcta y formatea la respuesta sin LLM. El "
        "modo se decide en tiempo de invocación con un detector de placeholder de la "
        "API key."
    )

    add_heading(doc, "4.2. Vista de despliegue", level=2)
    add_para(
        doc,
        "El despliegue se orquesta con Docker Compose v2. Cada servicio corre en su "
        "propio contenedor, comparte la red interna `smartsack_net` y, en el caso de "
        "PostgreSQL y Redis, persiste su estado en volúmenes nombrados. Esta "
        "arquitectura cumple con el requisito RNF-10 (portabilidad) y RNF-11 "
        "(disponibilidad)."
    )
    add_table(
        doc,
        ["Servicio", "Imagen", "Puerto", "Healthcheck", "Volumen"],
        [
            ["postgres", "postgres:16-alpine", "5432", "pg_isready", "smartsack_postgres_data"],
            ["redis", "redis:7-alpine", "6379", "redis-cli ping", "smartsack_redis_data"],
            ["backend", "build local (smartsack-backend)", "8000", "GET /health", "bind ./backend → /app"],
            ["frontend", "build local (smartsack-frontend)", "5173", "Vite dev server", "bind ./frontend → /app"],
            ["nginx", "build local (smartsack-nginx)", "80", "—", "—"],
        ],
    )
    add_para(
        doc,
        "El reverse proxy Nginx unifica los servicios en el puerto 80 con las "
        "siguientes reglas de enrutamiento: /api/ → backend:8000/api/, /docs → "
        "backend:8000/docs, /redoc → backend:8000/redoc, /ws/ → backend:8000/ws/ con "
        "upgrade WebSocket, y todo lo demás → frontend:5173. Esta unificación elimina "
        "los problemas de CORS y simplifica la configuración del cliente, que solo "
        "necesita conocer la URL base http://localhost."
    )

    add_heading(doc, "4.3. Vista de seguridad", level=2)
    add_para(
        doc,
        "El control de acceso se implementa mediante JWT (JSON Web Tokens). El flujo es "
        "el siguiente: el cliente envía credenciales a POST /api/auth/login con "
        "form-data (OAuth2PasswordRequestForm); el backend valida la contraseña con "
        "passlib/bcrypt; si es correcta, emite un JWT firmado con HMAC-SHA256 que "
        "codifica como claims `sub`, `uid`, `role`, `machine_id`, `iat`, `exp`. El "
        "token tiene una vigencia configurable (por defecto 480 minutos = 8 horas, "
        "equivalente a un turno de planta)."
    )
    add_para(
        doc,
        "Las dependencias de FastAPI implementan tres niveles de autorización: "
        "`get_current_user` (cualquier usuario autenticado), `require_supervisor` "
        "(supervisor o admin) y `require_admin` (solo admin). Cada router declara "
        "explícitamente su nivel mínimo. El endpoint WebSocket /ws/plant valida el "
        "token recibido como query parameter y cierra la conexión con código 1008 "
        "(Policy Violation) si es inválido o si el rol no es supervisor/admin."
    )
    add_para(
        doc,
        "Los secretos (JWT_SECRET_KEY, ANTHROPIC_API_KEY, DATABASE_URL) se inyectan "
        "exclusivamente vía variables de entorno. El archivo .env está "
        "explícitamente excluido en .gitignore (RNF-05). Para producción se recomienda "
        "delegar la gestión a Docker secrets, AWS Secrets Manager o HashiCorp Vault, "
        "como se documenta en el README."
    )

    add_heading(doc, "4.4. Decisiones arquitectónicas (ADR resumidos)", level=2)
    add_para(
        doc,
        "Las siguientes decisiones de arquitectura se adoptaron durante las fases 1 y "
        "2 del proyecto. Se documentan aquí en formato ADR (Architecture Decision "
        "Record) abreviado: contexto, decisión, consecuencias."
    )
    add_table(
        doc,
        ["ID", "Decisión", "Justificación", "Consecuencias"],
        [
            ["ADR-01", "Backend monolítico en Python (FastAPI) en lugar de microservicios", "Simplicidad de despliegue, menor sobrecarga operativa para un trabajo de grado, productividad del estudiante.", "Acoplamiento entre módulos. Mitigación: separación clara routers/services/models y dependencias explícitas en FastAPI."],
            ["ADR-02", "Postgres como única fuente de verdad, Redis solo como caché", "Postgres da consistencia transaccional para operaciones críticas (órdenes, eventos, OEE). Redis se usa para reducir round-trips en lectura de estado de máquinas.", "El caché puede quedar desincronizado; se acepta dado que su TTL es corto."],
            ["ADR-03", "WebSocket nativo de FastAPI/Starlette en lugar de SSE o polling", "Bidireccional, baja latencia, soportado por todos los navegadores objetivo.", "Requiere Nginx con `proxy_set_header Upgrade $http_upgrade`. Implementado en nginx.conf."],
            ["ADR-04", "Vista operario y supervisor en la misma SPA con rutas protegidas", "Un solo bundle, gestión común de auth, deploy unificado.", "Bundle ligeramente más pesado para el operario; aceptable dado el LAN local."],
            ["ADR-05", "ML embebido en el backend, no en microservicio separado", "Latencia mínima al servir predicciones; el modelo cabe en RAM (< 10 MB).", "Si en el futuro el modelo crece, se podría externalizar a TorchServe o BentoML."],
            ["ADR-06", "Function calling con LangChain en lugar de cliente Anthropic directo", "LangChain abstrae el ciclo tool→result→tool y simplifica añadir nuevas tools.", "Dependencia adicional. Aceptable: LangChain es estándar de facto."],
            ["ADR-07", "Fallback heurístico cuando no hay API key", "Permite demos sin credenciales y degradación elegante en cortes de Anthropic.", "Mantener dos rutas de respuesta. Mitigado por tests dedicados al fallback."],
            ["ADR-08", "JavaScript en lugar de TypeScript en el frontend", "Productividad para un trabajo de grado de un solo desarrollador; menor configuración.", "Pérdida de tipado estático. Se mitiga con PropTypes y tests en componentes críticos."],
            ["ADR-09", "Tailwind CSS en lugar de CSS-in-JS o frameworks UI completos", "Iteración rápida en diseño, sin lock-in, bundle pequeño.", "Requiere disciplina para no duplicar utilidades; se usan componentes comunes en components/common/."],
            ["ADR-10", "Pytest con sesión transaccional rollback como aislamiento", "No se necesita una BD separada para tests; el rollback garantiza independencia.", "Tests acoplados a la presencia del seed; los tests validan el catálogo del seed pero toleran extras (refactor del bloque 167 verde)."],
        ],
    )
    add_pagebreak(doc)


# -----------------------------------------------------------------------------
# 5. MODELO DE DATOS
# -----------------------------------------------------------------------------
def write_modelo_datos(doc):
    add_heading(doc, "5. MODELO DE DATOS", level=1)
    add_para(
        doc,
        "El modelo de datos relacional está implementado en PostgreSQL 16 y se "
        "materializa mediante SQLAlchemy 2.0 con la sintaxis declarativa moderna "
        "(Mapped/mapped_column). Las migraciones se gestionan con Alembic y el estado "
        "actual está en la revisión `22d436dd1c32`. La base de datos contiene catorce "
        "tablas de negocio más la tabla técnica `alembic_version`."
    )

    add_heading(doc, "5.1. Diagrama entidad-relación (descripción)", level=2)
    add_para(
        doc,
        "El diagrama ER de SmartSack se centra en cuatro entidades fuertes "
        "(Machine, ProductionOrder, OrderOperation, ProductionEvent) que articulan el "
        "flujo de producción, rodeadas de seis entidades de soporte (User, Shift, "
        "Material, OEERecord, MLPrediction, QualityRecord) y cuatro entidades "
        "técnicas/auditoría (ETLLoad, AdminAuditLog, SystemSetting, "
        "alembic_version)."
    )
    add_para(
        doc,
        "Las relaciones cardinales más relevantes son: "
        "una orden de producción se ejecuta en N operaciones consecutivas (1:N "
        "ProductionOrder → OrderOperation); cada operación se asigna a una máquina y, "
        "opcionalmente, a un operario (N:1 OrderOperation → Machine, OrderOperation → "
        "User); cada máquina puede recibir N eventos de producción (1:N Machine → "
        "ProductionEvent); cada orden tiene N predicciones del modelo ML a lo largo "
        "del tiempo (1:N ProductionOrder → MLPrediction); cada (máquina, turno, fecha) "
        "tiene un único registro de OEE (clave compuesta lógica en OEERecord). El "
        "diagrama ER recomendado se incluirá como imagen en el Anexo A; su "
        "especificación textual permite reproducirlo en draw.io o pgModeler."
    )

    add_heading(doc, "5.2. Inventario de tablas", level=2)
    add_table(
        doc,
        ["Tabla", "Campos", "Propósito"],
        [
            ["users", "8", "Usuarios del sistema (admin, supervisor, operario) con su rol y máquina asignada."],
            ["machines", "7", "Catálogo de máquinas de la planta, su tipo, ubicación y estado actual."],
            ["shifts", "4", "Turnos de operación (turno_1: 06–14, turno_2: 14–22, turno_3: 22–06)."],
            ["production_orders", "16", "Órdenes de producción con cantidades planeadas, real, fechas y estado."],
            ["order_operations", "16", "Operaciones individuales de cada orden, una por máquina-etapa."],
            ["production_events", "11", "Eventos en tiempo real registrados por operarios (paradas, reportes, etc.)."],
            ["materials", "6", "Materiales planeados/consumidos por orden (papel kraft, tinta, adhesivos)."],
            ["quality_records", "6", "Registros de calidad: producidas vs. defectuosas para cálculo de OEE.Q."],
            ["oee_records", "8", "Registros diarios de OEE por máquina y turno, con A/P/Q descompuesto."],
            ["ml_predictions", "7", "Predicciones del modelo XGBoost para cada orden activa."],
            ["etl_loads", "13", "Bitácora de cargas ETL (CSV→PostgreSQL) con métricas de validación."],
            ["admin_audit_log", "10", "Auditoría de acciones administrativas con diff antes/después."],
            ["system_settings", "5", "Configuraciones modificables en runtime (umbrales, etc.)."],
            ["alembic_version", "1", "Tabla técnica de Alembic para tracking de migraciones."],
        ],
    )

    add_heading(doc, "5.3. Diccionario de datos (campos clave)", level=2)
    add_para(
        doc,
        "El diccionario de datos completo, tabla por tabla y campo por campo, se "
        "encuentra en el Anexo A. Aquí se detallan únicamente las cuatro entidades "
        "centrales por su impacto en la lógica de negocio."
    )
    add_para(doc, "machines", bold=True)
    add_table(
        doc,
        ["Campo", "Tipo", "Restricciones", "Descripción"],
        [
            ["id", "INTEGER", "PK, autoincrement", "Identificador único."],
            ["code", "VARCHAR(32)", "UNIQUE, NOT NULL", "Código corto (TUB-01, IMP-02, EMP-01, etc.)."],
            ["name", "VARCHAR(64)", "NOT NULL", "Nombre legible para humanos."],
            ["type", "ENUM", "NOT NULL", "tubuladora | impresora | fondadora | empacadora."],
            ["location", "VARCHAR(64)", "NULL", "Ubicación física en la planta."],
            ["status", "ENUM", "NOT NULL", "running | idle | maintenance | down."],
            ["current_order_id", "INTEGER", "FK → production_orders.id, NULL", "Orden activa en la máquina, si la hay."],
        ],
    )
    add_para(doc, "production_orders", bold=True)
    add_table(
        doc,
        ["Campo", "Tipo", "Restricciones", "Descripción"],
        [
            ["id", "INTEGER", "PK", "Identificador único."],
            ["order_number", "VARCHAR(32)", "UNIQUE, NOT NULL", "Código de la orden (OP-2026-NNNNNN)."],
            ["product_type", "VARCHAR(64)", "NOT NULL, indexed", "Tipo de saco (Saco cemento 50kg, Saco harina 25kg, etc.)."],
            ["product_description", "VARCHAR(255)", "NULL", "Descripción libre del producto."],
            ["quantity_ordered", "INTEGER", "NOT NULL", "Sacos planeados."],
            ["quantity_produced", "INTEGER", "NOT NULL, default 0", "Sacos producidos hasta el momento."],
            ["scrap_total_kg", "FLOAT", "NOT NULL, default 0.0", "Acumulado de scrap en kg."],
            ["unit_weight_kg", "FLOAT", "NOT NULL, default 0.1", "Peso unitario para conversión sacos↔kg."],
            ["machine_id", "INTEGER", "FK → machines.id, NULL", "Máquina actualmente asignada."],
            ["status", "ENUM", "NOT NULL", "pending | in_progress | completed | delayed | cancelled."],
            ["priority", "ENUM", "NOT NULL", "low | normal | high | urgent."],
            ["planned_start", "TIMESTAMPTZ", "NOT NULL", "Inicio planeado."],
            ["planned_end", "TIMESTAMPTZ", "NOT NULL", "Fin planeado."],
            ["actual_start", "TIMESTAMPTZ", "NULL", "Inicio real."],
            ["actual_end", "TIMESTAMPTZ", "NULL", "Fin real."],
            ["created_at", "TIMESTAMPTZ", "NOT NULL", "Fecha de creación del registro."],
        ],
    )
    add_para(doc, "order_operations", bold=True)
    add_para(
        doc,
        "Una orden de producción se descompone en N operaciones secuenciales, una por "
        "etapa de la línea (Tubulado → Impresión → Fondado → Empacado). Cada "
        "operación puede tener un operario asignado, registra cantidades de entrada y "
        "salida (para calcular yield), kg de scrap con su razón, y se mueve por los "
        "estados pending → ready → in_progress → completed."
    )
    add_para(doc, "production_events", bold=True)
    add_para(
        doc,
        "Captura todo lo que sucede en tiempo real en una máquina: inicio de "
        "operación, parada planificada/no planificada, cambio de formato, incidencia, "
        "reporte parcial, fin de operación. Cada evento referencia opcionalmente a la "
        "orden y a la operación, y siempre lleva el operario que lo registró. Es la "
        "fuente principal del cálculo de Disponibilidad del OEE."
    )

    add_heading(doc, "5.4. Tipos enumerados", level=2)
    add_para(
        doc,
        "Los siguientes tipos enumerados, definidos en `app/models/enums.py`, garantizan "
        "integridad referencial a nivel de aplicación y se materializan como tipos "
        "ENUM nativos de PostgreSQL."
    )
    add_table(
        doc,
        ["Enum", "Valores"],
        [
            ["UserRole", "admin, supervisor, operator"],
            ["MachineType", "tubuladora, impresora, fondadora, empacadora"],
            ["MachineStatus", "running, idle, maintenance, down"],
            ["OrderStatus", "pending, in_progress, completed, delayed, cancelled"],
            ["OrderPriority", "low, normal, high, urgent"],
            ["OperationStatus", "pending, ready, in_progress, completed, blocked"],
            ["EventType", "start_operation, stop, format_change, incident, partial_report, end_operation"],
            ["ShiftName", "turno_1, turno_2, turno_3"],
            ["ScrapReason", "quality_defect, setup_loss, material_break, other"],
            ["ETLLoadKind", "production_orders, confirmations, materials, shipments"],
            ["ETLLoadStatus", "running, success, partial, failed"],
        ],
    )

    add_heading(doc, "5.5. Estrategia de migraciones", level=2)
    add_para(
        doc,
        "Las migraciones se gestionan con Alembic. La política adoptada es: "
        "(1) cada cambio de modelo se acompaña de una migración generada con "
        "`alembic revision --autogenerate`; (2) las migraciones se revisan manualmente "
        "antes de aplicarlas; (3) la rama de migración es lineal (sin merges); (4) en "
        "producción se aplica únicamente vía `alembic upgrade head` ejecutado durante "
        "el despliegue. Esta política previene divergencias y facilita rollbacks."
    )
    add_pagebreak(doc)


# -----------------------------------------------------------------------------
# 6. DISEÑO DE INTERFAZ
# -----------------------------------------------------------------------------
def write_interfaz(doc):
    add_heading(doc, "6. DISEÑO DE LA INTERFAZ DE USUARIO", level=1)
    add_para(
        doc,
        "El frontend de SmartSack es una Single Page Application (SPA) construida con "
        "React 18, JavaScript ES6+, Vite, Tailwind CSS y Recharts. La aplicación se "
        "organiza en ocho vistas principales y aproximadamente cuarenta componentes "
        "reutilizables agrupados por dominio. Cada vista responde a uno o más casos de "
        "uso de la sección 3 y aplica el contrato de autorización del JWT."
    )

    add_heading(doc, "6.1. Principios de diseño", level=2)
    add_bullets(
        doc,
        [
            "Diseño responsivo con prioridad al monitor estándar de PCs de planta (1366×768).",
            "Esquema de color pensado para reducir fatiga visual en jornadas largas: fondos claros, acentos en azul institucional, semáforos rojo/amarillo/verde para estados.",
            "Tipografía sans-serif (Inter por defecto) con tamaños generosos para legibilidad a un metro de distancia.",
            "Jerarquía visual clara: el dato principal es visualmente dominante, los metadatos están subordinados.",
            "Cero pantallas sin estado vacío; toda lista vacía muestra una ilustración o mensaje explicativo en español.",
            "Degradación elegante: si el chatbot falla, se muestra mensaje en español y se cae a fallback; si los WS pierden conexión, se hace polling cada 30s.",
        ],
    )

    add_heading(doc, "6.2. Inventario de vistas", level=2)
    add_table(
        doc,
        ["Vista", "Ruta", "Rol mínimo", "Caso de uso"],
        [
            ["Login", "/login", "Público", "CU-01"],
            ["Dashboard", "/dashboard", "Supervisor", "CU-09"],
            ["OperatorView", "/operator", "Operario", "CU-04, CU-05, CU-06, CU-07"],
            ["SupervisorView", "/supervisor", "Supervisor", "CU-08"],
            ["OrderTraceView", "/orders/:id/trace", "Supervisor", "CU-13 (trazabilidad)"],
            ["Chat", "/chat", "Supervisor", "CU-11"],
            ["ETL", "/etl", "Admin", "CU-12"],
            ["Admin", "/admin", "Admin", "CU-13, CU-14, CU-15"],
        ],
    )

    add_heading(doc, "6.3. Descripción de vistas principales", level=2)

    add_para(doc, "6.3.1. Login", bold=True)
    add_para(
        doc,
        "Vista pública con formulario de usuario y contraseña. En éxito, redirige según "
        "el rol: operario → /operator, supervisor → /dashboard, admin → /admin. Toda "
        "respuesta de error se muestra en español por debajo del formulario "
        "(\"Credenciales inválidas\", \"Cuenta inactiva\", etc.). El componente "
        "AuthContext almacena el token en localStorage y lo inyecta automáticamente en "
        "los headers Authorization de las llamadas a la API mediante un interceptor "
        "axios."
    )

    add_para(doc, "6.3.2. OperatorView (Vista de operario)", bold=True)
    add_para(
        doc,
        "Esta es la vista crítica para la captura en tiempo real. Diseñada para ser "
        "operada con el ratón (los PCs de planta no tienen pantalla táctil), muestra: "
        "(a) MachineStatusCard con el estado actual de la máquina del operario y un "
        "reloj en vivo; (b) CurrentOperationCard con la operación que está corriendo, "
        "su progreso en porcentaje y barra visual; (c) MachineActionButtons con cinco "
        "botones grandes para iniciar operación, registrar parada, cambio de formato, "
        "incidencia o reportar avance; (d) OperationQueue con las operaciones siguientes "
        "en cola; (e) RecentEventsLog con los últimos eventos de la máquina. Cada acción "
        "abre un Modal con el formulario mínimo requerido y se valida del lado cliente "
        "antes del envío."
    )

    add_para(doc, "6.3.3. SupervisorView (Digital Twin)", bold=True)
    add_para(
        doc,
        "Replica visualmente el piso de planta. PlantMap renderiza un grid de "
        "MachineTile, una por cada máquina presente en la base de datos. Cada tile "
        "muestra el código, nombre, estado (color de fondo según el estado: verde "
        "running, amarillo idle, rojo down, gris maintenance), la orden en curso, el "
        "progreso y los conteos de paradas/incidencias del día. La vista se conecta al "
        "canal /ws/plant; cuando llega un mensaje del tipo `state_change` o `event`, "
        "el tile correspondiente se actualiza sin recargar la página. El componente "
        "EventTicker muestra una banda inferior con el feed de eventos en tiempo real, "
        "y PlantStats agrega las cifras totales de la planta."
    )

    add_para(doc, "6.3.4. Dashboard de KPIs", bold=True)
    add_para(
        doc,
        "Vista agregada para análisis. Compone nueve widgets en una rejilla "
        "responsiva: KpiOverview (tarjetas con OEE, A/P/Q, contadores de órdenes, "
        "producción), OEEBreakdown (descomposición Disponibilidad·Rendimiento·Calidad), "
        "OEETrendChart (línea de tendencia 30 días con Recharts), MachineRanking "
        "(barras horizontales por máquina), AlertsPanel (lista de órdenes con "
        "probabilidad de retraso ≥ umbral), ProductionByShiftChart (barras apiladas "
        "por turno), OrderFulfillmentChart (cumplimiento de órdenes en el tiempo), "
        "ScrapPareto (top máquinas por kg de scrap) y WIPSnapshot (work in progress "
        "por máquina). Todos los widgets aceptan filtros de fecha mediante "
        "controladores de URL (`?days=7`)."
    )

    add_para(doc, "6.3.5. Chat conversacional", bold=True)
    add_para(
        doc,
        "Interfaz tipo mensajería. El usuario escribe en lenguaje natural y recibe la "
        "respuesta del asistente. La UI muestra los `tool_calls` que ejecutó el "
        "modelo, lo que añade transparencia: el usuario ve qué función se invocó y "
        "con qué argumentos. Si el modo es `fallback`, una etiqueta lo indica de "
        "manera explícita para no engañar al usuario. El historial se mantiene "
        "client-side y se envía completo en cada solicitud (stateless server-side)."
    )

    add_para(doc, "6.3.6. ETL", bold=True)
    add_para(
        doc,
        "Vista para administradores. UploadDropzone permite arrastrar uno o varios "
        "archivos CSV; cada uno se identifica por tipo desde un dropdown. El sistema "
        "muestra el progreso, el resultado del ETL en una tarjeta resumen "
        "(insertados/actualizados/omitidos/fallidos/duración) y mantiene un "
        "LoadHistory con las cargas previas. Errores específicos por fila se "
        "exponen en un modal con la línea problemática."
    )

    add_para(doc, "6.3.7. Admin", bold=True)
    add_para(
        doc,
        "Panel administrativo dividido en cuatro pestañas: AdminUsers (alta, edición, "
        "asignación de máquina, reset de contraseña, desactivación), AdminMachines "
        "(CRUD de máquinas), AdminOrders (creación/modificación de órdenes para "
        "demos), AdminAudit (consulta filtrada de la auditoría) y AdminSystem "
        "(estado de PostgreSQL, Redis, Anthropic, modelo ML; recarga del modelo)."
    )

    add_heading(doc, "6.4. Componentes reutilizables", level=2)
    add_para(
        doc,
        "El proyecto cuenta con un sistema de componentes reutilizables organizado por "
        "dominio en `frontend/src/components/`."
    )
    add_table(
        doc,
        ["Carpeta", "Componentes principales", "Función"],
        [
            ["common/", "Badge, Button, Card, Modal, Spinner, StatusDot, LiveClock", "Building blocks de UI usados en todas las vistas."],
            ["layout/", "Layout, Navbar, Sidebar, ProtectedRoute", "Estructura general de la app, navegación y protección por rol."],
            ["operator/", "MachineStatusCard, CurrentOperationCard, MachineActionButtons, OperationQueue, RecentEventsLog", "Componentes específicos de la vista de operario."],
            ["supervisor/", "MachineTile, PlantMap, EventTicker, PlantStats", "Componentes del Digital Twin del supervisor."],
            ["dashboard/", "KpiOverview, OEEBreakdown, OEETrendChart, MachineRanking, AlertsPanel, ProductionByShiftChart, OrderFulfillmentChart, ScrapPareto, WIPSnapshot, ModelInsights", "Widgets de analítica y KPIs (todos con Recharts)."],
            ["etl/", "UploadDropzone, LoadHistory", "Carga de archivos y bitácora de ETL."],
            ["admin/", "AdminUsers, AdminMachines, AdminOrders, AdminAudit, AdminSystem, Field", "Pestañas de gestión administrativa."],
            ["chat/", "(componentes embebidos en la página Chat)", "Mensajería del asistente conversacional."],
        ],
    )

    add_heading(doc, "6.5. Mockups y capturas", level=2)
    add_para(
        doc,
        "Las capturas reales del sistema funcionando se incluirán en el Anexo D. La "
        "metodología recomendada para producirlas es: (1) levantar el sistema con "
        "`docker compose up -d`; (2) acceder a http://localhost; (3) iniciar sesión "
        "con cada uno de los tres roles (admin, supervisor, operario); (4) capturar "
        "cada vista en estado pleno (no estado vacío). Para cumplir con la exigencia "
        "del jurado, se recomienda incluir al menos doce capturas: Login, Dashboard "
        "completo, Dashboard con filtros, OperatorView con operación activa, "
        "SupervisorView con planta llena, OrderTraceView, Chat con varios mensajes, "
        "ETL antes y después de subir archivo, Admin (cada pestaña)."
    )
    add_pagebreak(doc)


# -----------------------------------------------------------------------------
# 7. DISEÑO DEL MODELO ML
# -----------------------------------------------------------------------------
def write_ml(doc):
    add_heading(doc, "7. DISEÑO DEL MODELO DE MACHINE LEARNING", level=1)
    add_para(
        doc,
        "El módulo de Machine Learning de SmartSack se diseñó siguiendo estrictamente "
        "el ciclo CRISP-DM (Cross-Industry Standard Process for Data Mining). La "
        "presente sección documenta cada fase del ciclo aplicada al objetivo concreto "
        "del proyecto: predecir la probabilidad de retraso de una orden de producción "
        "antes de que ocurra. Las métricas reportadas corresponden al entrenamiento "
        "actual sobre datos sintéticos generados por el script de seed; en la Entrega 3 "
        "se reentrenará con datos reales del ERP y se reportarán las métricas "
        "definitivas para validar la hipótesis H2 (F1 ≥ 0.80)."
    )

    add_heading(doc, "7.1. Comprensión del negocio", level=2)
    add_para(
        doc,
        "El problema de negocio es la detección anticipada de órdenes en riesgo de "
        "retraso, definido formalmente como una clasificación binaria: dado un vector "
        "de features de una orden, predecir si su tiempo real de finalización superará "
        "el tiempo planeado en más de un margen tolerado. La variable objetivo "
        "`is_delayed` se computa como `actual_end > planned_end` para órdenes "
        "completadas; durante el entrenamiento se descartan órdenes canceladas o "
        "incompletas. La importancia de adelantar esta señal radica en habilitar "
        "intervenciones operativas (reprogramar, reasignar, comunicar al cliente) que "
        "transformen un retraso evitable en una entrega a tiempo."
    )

    add_heading(doc, "7.2. Comprensión de los datos", level=2)
    add_para(
        doc,
        "Los datos provienen de tres fuentes: (1) órdenes históricas cargadas vía ETL "
        "desde el ERP en formato CSV; (2) eventos de producción registrados por los "
        "operarios en tiempo real; (3) catálogos maestros (máquinas, productos, "
        "operarios). En el entorno actual estos datos están generados por un script de "
        "seed que sintetiza seis meses de actividad de planta, garantizando volumetría "
        "suficiente para entrenar y validar."
    )
    add_table(
        doc,
        ["Métrica del dataset", "Valor"],
        [
            ["Filas totales (órdenes con label)", "3,331"],
            ["Positivos (retraso confirmado)", "483"],
            ["Tasa de positivos", "14.50%"],
            ["Filas en test set (test_size = 0.2)", "667"],
            ["Positivos en test set", "97"],
            ["Número de features tras one-hot", "30"],
        ],
    )
    add_para(
        doc,
        "El desbalance de clases (14.5% positivos) es típico en problemas operativos: "
        "la mayoría de las órdenes salen a tiempo. Este desbalance se gestiona mediante "
        "(a) `class_weight='balanced'` en Random Forest, (b) `scale_pos_weight=5` en "
        "XGBoost, y (c) métricas robustas a desbalance (F1, AUC-ROC) en lugar de "
        "accuracy."
    )

    add_heading(doc, "7.3. Preparación de los datos y feature engineering", level=2)
    add_para(
        doc,
        "El pipeline de preparación está implementado en `backend/ml/features.py` y "
        "produce un DataFrame con treinta columnas tras la codificación one-hot. Los "
        "features se agrupan en cinco categorías:"
    )
    add_para(doc, "Numéricas crudas (4 features)", bold=True)
    add_bullets(
        doc,
        [
            "quantity_ordered: cantidad planeada en unidades.",
            "planned_duration_hours: diferencia (planned_end − planned_start) en horas.",
            "hour_of_day: hora del día del inicio planeado (0–23).",
            "day_of_week: día de la semana del inicio planeado (0–6).",
        ],
    )
    add_para(doc, "Indicadores binarios (1 feature)", bold=True)
    add_bullets(doc, ["is_weekend: 1 si day_of_week ∈ {5, 6}, 0 en caso contrario."])
    add_para(doc, "Features derivados (3 features)", bold=True)
    add_bullets(
        doc,
        [
            "machine_concurrent_load: número de órdenes activas en la misma máquina al inicio planeado.",
            "machine_delay_rate_30d: tasa de retrasos de la máquina en los últimos 30 días.",
            "product_delay_rate_30d: tasa de retrasos del tipo de producto en los últimos 30 días.",
        ],
    )
    add_para(doc, "Categóricos one-hot (22 features)", bold=True)
    add_bullets(
        doc,
        [
            "product_type: 6 columnas (Saco cemento 50/25kg, Saco cal 25kg, Saco fertilizante 25kg, Saco harina 50kg, _UNK).",
            "machine_code: 9 columnas (TUB-01, TUB-02, IMP-01, IMP-02, FON-01, FON-02, EMP-01, EMP-02, _UNK).",
            "shift: 3 columnas (turno_1, turno_2, turno_3).",
            "priority: 4 columnas (low, normal, high, urgent).",
        ],
    )
    add_para(
        doc,
        "La codificación one-hot incluye una columna `_UNK` para valores no vistos "
        "durante el entrenamiento, lo que aporta robustez ante nuevos productos o "
        "máquinas. Los features derivados (`machine_delay_rate_30d`, "
        "`product_delay_rate_30d`) se calculan con una ventana móvil que respeta la "
        "fecha de la orden a clasificar, evitando data leakage."
    )

    add_heading(doc, "7.4. Modelado", level=2)
    add_para(
        doc,
        "Se evaluaron dos algoritmos de gradient boosting / ensembling de árboles, "
        "elegidos por su buen rendimiento en datos tabulares estructurados: Random "
        "Forest (Scikit-learn) y XGBoost. Para cada uno se ejecutó búsqueda de "
        "hiperparámetros con GridSearchCV y validación cruzada estratificada de 5 "
        "folds, optimizando F1-score. La estratificación preserva la proporción de "
        "positivos en cada fold, lo cual es crítico para datos desbalanceados."
    )
    add_para(
        doc,
        "El espacio de búsqueda de hiperparámetros para Random Forest fue: "
        "n_estimators ∈ {80, 120, 200}, max_depth ∈ {6, 12, None}, "
        "min_samples_leaf ∈ {1, 2, 4}, class_weight ∈ {None, 'balanced'}. Para "
        "XGBoost: n_estimators ∈ {100, 200}, max_depth ∈ {3, 4, 6}, "
        "learning_rate ∈ {0.05, 0.1}, subsample ∈ {0.8, 0.9, 1.0}. Se utilizó "
        "scoring='f1' y refit en el mejor F1 promedio de los folds."
    )

    add_heading(doc, "7.5. Evaluación", level=2)
    add_para(
        doc,
        "La tabla siguiente reporta las métricas obtenidas sobre el test set "
        "(20% holdout) por cada algoritmo. El modelo ganador se determinó por F1-score "
        "en cross-validation."
    )
    add_table(
        doc,
        ["Métrica", "Random Forest", "XGBoost (ganador)"],
        [
            ["F1 (CV best)", "0.0437", "0.1798"],
            ["F1 (test)", "0.0367", "0.2110"],
            ["Precision (test)", "0.1667", "0.1786"],
            ["Recall (test)", "0.0206", "0.2577"],
            ["AUC-ROC (test)", "0.4772", "0.5260"],
            ["Tiempo de entrenamiento", "13.9 s", "11.45 s"],
            ["Mejores hiperparámetros", "n=120, depth=12, leaf=2, balanced", "n=200, depth=4, lr=0.1, subsample=0.9"],
        ],
    )
    add_para(
        doc,
        "Las métricas actuales se ubican significativamente por debajo del objetivo "
        "H2 (F1 ≥ 0.80) declarado en la Entrega 1. Esta brecha es consistente con la "
        "naturaleza sintética del dataset: el script de seed introduce variabilidad "
        "aleatoria sin patrones causales fuertes entre los features y el label, lo que "
        "limita la información explotable por cualquier algoritmo. El AUC-ROC cercano "
        "a 0.5 confirma que, sobre datos puramente aleatorios, el modelo apenas "
        "supera el desempeño de un clasificador trivial."
    )
    add_para(
        doc,
        "El objetivo del entrenamiento sobre datos sintéticos no es alcanzar las "
        "métricas finales sino validar que el pipeline completo funciona "
        "extremo-a-extremo: extracción de features, codificación, entrenamiento, "
        "validación cruzada, persistencia con joblib, carga en runtime y predicción "
        "vía endpoint REST. Todo este flujo está operativo y se demostrará en la "
        "Entrega 3 con datos reales del ERP, momento en el cual se espera que las "
        "tasas de retraso reales (correlacionadas con máquina, producto, prioridad y "
        "carga concurrente) eleven el F1 por encima del objetivo."
    )
    add_para(doc, "7.5.1. Top-15 feature importance (XGBoost)", bold=True)
    add_para(
        doc,
        "El análisis de feature importance del modelo XGBoost actual muestra una "
        "distribución relativamente uniforme, sin una variable dominante. Esto es "
        "esperable en datos sintéticos donde no hay relaciones causales fuertes; con "
        "datos reales se anticipa que `machine_delay_rate_30d`, `product_delay_rate_30d` "
        "y `machine_concurrent_load` concentren la importancia, dado que codifican "
        "evidencia operativa directa."
    )
    add_table(
        doc,
        ["Feature", "Importance"],
        [
            ["product_type__Saco harina 50kg", "0.0488"],
            ["priority__low", "0.0458"],
            ["machine_concurrent_load", "0.0447"],
            ["planned_duration_hours", "0.0438"],
            ["priority__normal", "0.0431"],
            ["product_type__Saco cemento 50kg", "0.0427"],
            ["product_type__Saco cemento 25kg", "0.0426"],
            ["machine_delay_rate_30d", "0.0424"],
            ["priority__urgent", "0.0400"],
            ["quantity_ordered", "0.0392"],
            ["product_delay_rate_30d", "0.0391"],
            ["day_of_week", "0.0380"],
            ["machine_code__TUB-01", "0.0379"],
            ["machine_code__EMP-01", "0.0378"],
            ["hour_of_day", "0.0374"],
        ],
    )

    add_heading(doc, "7.6. Despliegue del modelo", level=2)
    add_para(
        doc,
        "El modelo entrenado se persiste en `backend/ml/models/delay_predictor.joblib` "
        "junto a un manifest JSON (`delay_predictor.manifest.json`) que documenta "
        "versión, fecha, dataset, hiperparámetros, métricas e importancias. El servicio "
        "`prediction_service.py` carga el artefacto al arranque de FastAPI y lo expone "
        "vía cuatro endpoints: GET /api/predictions/model-info, GET /api/predictions/"
        "feature-importance, POST /api/predictions/predict/{order_id} y POST "
        "/api/predictions/predict-active. Cada predicción se persiste en "
        "`ml_predictions` para soportar análisis longitudinal del comportamiento del "
        "modelo en producción."
    )
    add_para(
        doc,
        "El reentrenamiento se ejecuta vía CLI: `docker compose exec backend python -m "
        "ml.train` regenera artefacto y manifest. La política definida es: "
        "reentrenar al cargar nuevos datos del ERP cuando la cantidad incremental "
        "supere el 10% del dataset previo, o cuando un monitoreo offline detecte "
        "drift en el AUC-ROC superior a 5 puntos porcentuales."
    )
    add_pagebreak(doc)


# -----------------------------------------------------------------------------
# 8. DISEÑO DEL ETL Y CAPTURA EN TIEMPO REAL
# -----------------------------------------------------------------------------
def write_etl(doc):
    add_heading(doc, "8. DISEÑO DEL ETL Y LA CAPTURA EN TIEMPO REAL", level=1)
    add_para(
        doc,
        "SmartSack adopta una arquitectura de doble fuente de datos para reflejar "
        "fielmente la realidad de la planta: un proceso ETL por lotes que importa "
        "exportaciones periódicas del ERP, y una capa de captura en tiempo real que "
        "los operarios usan desde el PC de su máquina. Ambas fuentes alimentan la "
        "misma base de datos PostgreSQL, manteniendo un único punto de verdad."
    )

    add_heading(doc, "8.1. Proceso ETL desde CSV del ERP", level=2)
    add_para(
        doc,
        "El proceso ETL acepta cuatro tipos de archivo: `production_orders` (órdenes "
        "del ERP), `confirmations` (confirmaciones de fabricación con cantidades "
        "reales), `materials` (materiales planeados/consumidos) y `shipments` "
        "(despachos). Cada tipo tiene un esquema definido (columnas obligatorias, "
        "tipos esperados, restricciones) que se valida antes de procesar. El servicio "
        "`etl_service.py` implementa el pipeline con Pandas siguiendo cinco etapas: "
        "(1) lectura del CSV con detección automática de encoding y separador; "
        "(2) validación de cabeceras contra el esquema esperado; (3) coerción de "
        "tipos y validaciones por fila (no nulos en campos obligatorios, formato de "
        "fechas, rangos numéricos); (4) detección de duplicados contra los registros "
        "ya existentes en la BD; (5) inserción/actualización transaccional con "
        "registro detallado en `etl_loads`."
    )
    add_para(
        doc,
        "La carga es idempotente: ejecutar dos veces el mismo CSV no duplica registros "
        "porque las claves naturales (order_number en orders, una tupla (order_id, "
        "operation_seq) en confirmations) detectan el duplicado y se actualizan en "
        "lugar de insertar. Las filas que fallan validación no detienen el proceso: "
        "se acumulan en un log de errores devuelto al cliente para su corrección "
        "manual. Esto garantiza que un CSV con 1000 filas y 5 problemáticas cargue "
        "las 995 válidas en lugar de fallar completamente."
    )
    add_para(
        doc,
        "La bitácora `etl_loads` registra para cada carga: filename, kind, status "
        "(running/success/partial/failed), uploaded_by_id, uploaded_at, rows_total, "
        "rows_inserted, rows_updated, rows_skipped, rows_failed, duration_ms y "
        "error_log (JSON). Esta bitácora es consultable por el administrador desde la "
        "vista ETL y queda como evidencia ante auditorías."
    )

    add_heading(doc, "8.2. Captura de eventos en tiempo real", level=2)
    add_para(
        doc,
        "Los operarios registran eventos desde la vista OperatorView. Cada acción "
        "(iniciar operación, registrar parada, cambio de formato, incidencia, reporte "
        "parcial, cerrar operación) genera una llamada a un endpoint específico de "
        "/api/operations o /api/events. El backend valida (autorización, máquina del "
        "operario coincide con la operación, transición de estados válida), persiste "
        "en `production_events` y actualiza el estado derivado (machine.status, "
        "operation.status, order.actual_start). Inmediatamente después, emite un "
        "broadcast por WebSocket a todos los supervisores conectados al canal "
        "/ws/plant para que sus tiles se actualicen sin recargar."
    )
    add_para(
        doc,
        "Los tipos de evento están enumerados en `EventType`: start_operation, "
        "stop, format_change, incident, partial_report, end_operation. Para "
        "eventos de scrap, el modelo OrderOperation lleva los campos `scrap_kg` y "
        "`scrap_reason` que alimentan directamente el dashboard de Pareto."
    )

    add_heading(doc, "8.3. Sincronización entre fuentes", level=2)
    add_para(
        doc,
        "La política de conciliación es: la fuente ERP es la verdad sobre `quantity_ordered`, "
        "`planned_start`, `planned_end`, `priority` y `product_type` (datos "
        "transaccionales del ERP); la fuente tiempo real es la verdad sobre "
        "`quantity_produced`, `actual_start`, `actual_end`, `scrap_total_kg` y "
        "`status` (datos operativos del piso). Cuando una orden recién importada del "
        "ERP entra en conflicto con eventos ya capturados en tiempo real, prevalecen "
        "los eventos de tiempo real para los campos operativos. Esta política se "
        "implementa en el servicio ETL con merge selectivo: solo se actualizan los "
        "campos del dominio del ERP, los del dominio operativo se conservan."
    )

    add_heading(doc, "8.4. Plantillas y descarga", level=2)
    add_para(
        doc,
        "Para reducir la fricción de la integración, SmartSack expone GET "
        "/api/etl/sample-csv/{kind} que devuelve un CSV vacío con las cabeceras "
        "exactas que espera cada tipo de carga. El administrador descarga la "
        "plantilla, la rellena (por copia desde el ERP), y la sube de vuelta. Este "
        "flujo elimina el riesgo de errores por nombres de columna o formato."
    )
    add_pagebreak(doc)


# -----------------------------------------------------------------------------
# 9. DISEÑO DEL ASISTENTE CONVERSACIONAL
# -----------------------------------------------------------------------------
def write_chatbot(doc):
    add_heading(doc, "9. DISEÑO DEL ASISTENTE CONVERSACIONAL", level=1)
    add_para(
        doc,
        "El asistente conversacional de SmartSack permite consultar datos de producción "
        "en español natural. Es la materialización del Objetivo 5 (Entrega 1) y el "
        "tercer pilar tecnológico del proyecto. Se diseña para operar en dos modos "
        "complementarios: un modo LLM con la API de Claude y un modo fallback "
        "heurístico, ambos consumiendo el mismo conjunto de ocho herramientas."
    )

    add_heading(doc, "9.1. Arquitectura del chatbot", level=2)
    add_para(
        doc,
        "El módulo se organiza en tres capas: (1) las tools, en `app/chat/tools.py`, "
        "son funciones Python puras que ejecutan consultas SQL contra PostgreSQL y "
        "devuelven dicts JSON-serializables; (2) los schemas, también en tools.py, "
        "describen las herramientas en formato compatible con Anthropic tools/OpenAI "
        "function calling; (3) el servicio, en `app/services/chat_service.py`, "
        "orquesta la conversación, decide entre modo LLM o fallback, y formatea la "
        "respuesta final."
    )

    add_heading(doc, "9.2. Catálogo de herramientas", level=2)
    add_para(
        doc,
        "Las ocho herramientas implementadas cubren los principales escenarios de "
        "consulta operativa identificados en entrevistas y diseño:"
    )
    add_table(
        doc,
        ["#", "Herramienta", "Descripción", "Tabla(s) consultada(s)"],
        [
            ["1", "get_production_stats", "Sacos producidos, órdenes completadas, tasa de cumplimiento por ventana temporal y filtros opcionales (máquina, producto).", "production_orders"],
            ["2", "get_machine_status", "Estado actual de una o todas las máquinas con conteo de paradas/incidencias en la ventana.", "machines, production_events"],
            ["3", "get_order_info", "Detalle completo de una orden + última predicción ML.", "production_orders, machines, ml_predictions"],
            ["4", "get_oee_data", "OEE de planta o de una máquina con descomposición A/P/Q.", "oee_records, machines"],
            ["5", "get_alerts", "Órdenes activas con probabilidad de retraso ≥ umbral.", "production_orders, ml_predictions, machines"],
            ["6", "get_scrap_summary", "Pareto de scrap por máquina y razón.", "order_operations, machines"],
            ["7", "get_yield_summary", "Yield (out/in) por máquina, identificación de cuellos de botella de calidad.", "order_operations, machines"],
            ["8", "get_wip_status", "Snapshot de Work-In-Progress por máquina (operaciones e unidades en cola/en curso).", "order_operations, machines"],
        ],
    )

    add_heading(doc, "9.3. Modo LLM (Claude + LangChain)", level=2)
    add_para(
        doc,
        "Cuando hay una API key válida de Anthropic configurada, el servicio invoca "
        "ChatAnthropic (LangChain) con `bind_tools(TOOL_SCHEMAS)`. El modelo recibe "
        "el system prompt en español que lo identifica como SmartSack Assistant, lo "
        "instruye sobre formato (separador de miles, tono profesional, brevedad de 4–5 "
        "frases), define el contexto del dominio (8 máquinas, 3 turnos, productos "
        "típicos) y define reglas de operación (no inventar datos, redirigir consultas "
        "fuera del dominio, usar defaults razonables ante ambigüedad)."
    )
    add_para(
        doc,
        "El bucle de tool calling iterа hasta cuatro veces: el modelo decide qué "
        "herramienta invocar, el servicio la ejecuta, devuelve el resultado en un "
        "ToolMessage, el modelo recibe ese contexto y o decide invocar otra "
        "herramienta o produce la respuesta final en español. Si tras cuatro "
        "iteraciones no se llega a respuesta concluyente, el sistema responde "
        "elegantemente que la pregunta debe ser reformulada. El parámetro de "
        "temperatura es 0.2, lo que privilegia respuestas reproducibles sobre "
        "creatividad, y `max_tokens` está limitado a 1024 por respuesta."
    )

    add_heading(doc, "9.4. Modo fallback (heurístico sin LLM)", level=2)
    add_para(
        doc,
        "Cuando la API key no está configurada o el LLM falla con error de red o de "
        "autenticación, el servicio cae a un router heurístico por palabras clave "
        "definido en una lista ordenada de reglas. Cada regla mapea un conjunto de "
        "palabras del español ('alerta', 'oee', 'parada', 'wip', etc.) a una "
        "herramienta. Detectores adicionales extraen el código de máquina "
        "(TUB-01, IMP-2, etc.), el número de orden (OP-2026-001234) y la ventana "
        "temporal ('ayer', 'última semana', 'últimos 14 días') del mensaje. Con esa "
        "información se construyen los argumentos y se invoca la tool, formateando "
        "manualmente la respuesta en plantillas predefinidas."
    )
    add_para(
        doc,
        "El fallback no depende de la API de Claude y por tanto no incurre en costos "
        "ni latencia de red. Es la red de seguridad que garantiza que el chatbot "
        "responda siempre algo útil, incluso en una demo sin internet o con cuotas "
        "agotadas. Las pruebas automatizadas (test_chat.py) cubren extensivamente este "
        "modo con 31 tests."
    )

    add_heading(doc, "9.5. Mecanismos de control y observabilidad", level=2)
    add_bullets(
        doc,
        [
            "Detección de placeholder: la función `is_llm_available()` rechaza claves obvias como cadenas vacías, 'tu_api_key_de_anthropic_aqui', o cualquier valor que no comience por 'sk-ant-'.",
            "Limitación de historial: solo se envían los últimos 20 mensajes al LLM para acotar costos y evitar saturar la ventana de contexto.",
            "Validación de payload: el endpoint POST /api/chat/message valida con Pydantic que el mensaje tenga máximo 2000 caracteres y que cada mensaje del historial tenga rol válido (user/assistant).",
            "Auditoría: cada respuesta del chat reporta `mode` (llm/fallback) y `tool_calls` con argumentos y preview del resultado, lo que da total transparencia al usuario y permite depuración.",
            "Endpoint /api/chat/status: expone si el LLM está disponible y la lista de keywords del fallback, útil para que el frontend muestre un badge `Modo Claude` vs `Modo Fallback`.",
        ],
    )
    add_pagebreak(doc)


# -----------------------------------------------------------------------------
# 10. ESTRATEGIA DE PRUEBAS
# -----------------------------------------------------------------------------
def write_pruebas(doc):
    add_heading(doc, "10. ESTRATEGIA DE PRUEBAS", level=1)
    add_para(
        doc,
        "La calidad del software se garantiza mediante una pirámide de pruebas que "
        "combina pruebas unitarias, de integración y de sistema. La cobertura actual "
        "es de 197 pruebas automatizadas: 167 en backend con Pytest y 30 en frontend "
        "con Vitest. Esta sección describe la estrategia, las herramientas y la "
        "distribución actual."
    )

    add_heading(doc, "10.1. Niveles de prueba", level=2)
    add_para(doc, "Pruebas unitarias", bold=True)
    add_para(
        doc,
        "Verifican el comportamiento de funciones individuales aisladas de "
        "dependencias. En el backend, son las pruebas de helpers de detección "
        "(_detect_machine_code, _detect_days_back, _detect_order_number, "
        "_route_keywords), las pruebas de tools del chat sin pasar por HTTP, y los "
        "componentes individuales del frontend (Badge, Modal, Field, MachineTile, "
        "OperationQueue)."
    )
    add_para(doc, "Pruebas de integración", bold=True)
    add_para(
        doc,
        "Verifican la interacción entre módulos y con la base de datos real. Toda "
        "prueba que use `client` (TestClient de FastAPI) o `db_session` (Session "
        "transaccional con rollback) cae en este nivel. La estrategia de "
        "transacción + rollback (definida en `tests/conftest.py`) garantiza "
        "aislamiento entre tests sin requerir una base de datos separada y sin "
        "contaminar el entorno de desarrollo."
    )
    add_para(doc, "Pruebas de sistema", bold=True)
    add_para(
        doc,
        "Verifican el sistema completo orquestado por Docker Compose. Los smoke tests "
        "validados manualmente durante el desarrollo de esta entrega incluyen: "
        "GET /api/health responde 200, POST /api/auth/login devuelve token válido, "
        "GET /api/machines devuelve catálogo, POST /api/predictions/predict-active "
        "genera predicciones, WebSocket /ws/plant emite snapshot tras conectar."
    )

    add_heading(doc, "10.2. Cobertura actual del backend (Pytest)", level=2)
    add_table(
        doc,
        ["Archivo", "Tests", "Dominio cubierto"],
        [
            ["test_admin.py", "12", "Endpoints /api/admin: auditoría, configuración, ml-status."],
            ["test_auth.py", "7", "Login, logout, /me, JWT inválido."],
            ["test_chat.py", "31", "Tools, helpers de detección, fallback, router /api/chat."],
            ["test_dashboard.py", "10", "/api/dashboard: overview, oee-trend, alerts, ranking, etc."],
            ["test_etl.py", "17", "Carga CSV: validación de cabeceras, duplicados, errores por fila."],
            ["test_events.py", "6", "POST/GET /api/events con filtros."],
            ["test_health.py", "2", "/health y /api/health."],
            ["test_machines.py", "9", "CRUD de máquinas, autorización por rol."],
            ["test_operations.py", "7", "Iniciar/cerrar/reportar operación."],
            ["test_orders.py", "7", "Listado, filtrado y consulta de órdenes."],
            ["test_predictions.py", "17", "predict, predict-active, model-info, feature-importance."],
            ["test_users.py", "20", "CRUD de usuarios, asignación de máquina, reset password."],
            ["test_websocket_manager.py", "5", "Manager: connect/disconnect/broadcast."],
            ["TOTAL", "150", "Suma directa por archivo (167 totales con parametrize)."],
        ],
    )
    add_para(
        doc,
        "Adicionalmente al conteo por archivo, los tests parametrizados con "
        "`@pytest.mark.parametrize` multiplican casos: en particular "
        "`test_route_keywords` y `test_detect_*` recorren múltiples ejemplos de "
        "entrada/salida. El total efectivo de assertions ejecutadas supera las 200."
    )

    add_heading(doc, "10.3. Cobertura actual del frontend (Vitest)", level=2)
    add_table(
        doc,
        ["Archivo", "Tests", "Componente cubierto"],
        [
            ["src/context/tokenStorage.test.js", "4", "Almacenamiento del JWT en localStorage con codificación."],
            ["src/components/common/Badge.test.jsx", "4", "Render condicional según variant y children."],
            ["src/components/common/Modal.test.jsx", "4", "Apertura, cierre, foco y bloqueo de scroll."],
            ["src/components/admin/Field.test.jsx", "5", "Validación inline, mensajes de error en español."],
            ["src/components/admin/AdminUsers.test.jsx", "4", "Render de lista, paginación, acciones."],
            ["src/components/supervisor/MachineTile.test.jsx", "4", "Color por estado, render del current_order."],
            ["src/components/operator/OperationQueue.test.jsx", "5", "Orden de la cola, marcado del activo."],
            ["TOTAL", "30", ""],
        ],
    )

    add_heading(doc, "10.4. Datos de prueba", level=2)
    add_para(
        doc,
        "Las pruebas asumen la presencia del seed (`scripts/seed.py`) que genera 8 "
        "máquinas, 24 operarios, 3 turnos y aproximadamente seis meses de actividad "
        "(órdenes, operaciones, eventos, OEE). El seed es determinístico (random.seed "
        "fija) y se ejecuta una sola vez por entorno. Las pruebas han sido refactorizadas "
        "para tolerar máquinas extra añadidas en runtime (caso real: una IMP-03 fue "
        "creada vía la UI para una demo) — verifican que las 8 máquinas del seed estén "
        "presentes pero aceptan más, lo que las hace robustas al estado mutable del "
        "entorno."
    )

    add_heading(doc, "10.5. Continuous Integration (recomendado para E3)", level=2)
    add_para(
        doc,
        "Para la Entrega 3 se recomienda configurar GitHub Actions con un workflow "
        "que: (1) levante Postgres y Redis como servicios; (2) construya las imágenes "
        "del backend y frontend; (3) ejecute `pytest` y `npm run test`; (4) reporte "
        "los resultados como artefactos del PR. El esqueleto del workflow se incluye "
        "en el Anexo C como referencia."
    )
    add_pagebreak(doc)


# -----------------------------------------------------------------------------
# 11. AVANCE DE IMPLEMENTACIÓN
# -----------------------------------------------------------------------------
def write_avance(doc):
    add_heading(doc, "11. AVANCE DE IMPLEMENTACIÓN", level=1)
    add_para(
        doc,
        "Esta sección reporta el grado de avance verificable del prototipo a la fecha "
        "de cierre de la Entrega 2. La verificación se realizó en el entorno de "
        "desarrollo del estudiante, con los cinco contenedores corriendo y la base de "
        "datos cargada con seis meses de datos sintéticos. Cada objetivo específico de "
        "la Entrega 1 se contrasta con su evidencia técnica y se etiqueta como "
        "completo, parcial o pendiente."
    )

    add_heading(doc, "11.1. Estado por objetivo específico", level=2)
    add_table(
        doc,
        ["Objetivo (E1 §3.2)", "Estado", "Evidencia"],
        [
            ["O1. Diagnóstico y arquitectura", "Completo", "Docker Compose con 5 servicios, modelo de datos con 14 entidades, README, ADR (sección 4.4)."],
            ["O2. ETL del ERP a PostgreSQL + tiempo real", "Completo (sobre datos sintéticos)", "etl_service.py (802 LOC) maneja 4 tipos de CSV. Endpoint POST /api/etl/upload + bitácora en etl_loads. Captura tiempo real vía /api/events y /api/operations."],
            ["O3. Digital Twin con WebSockets", "Completo", "8 vistas frontend, WebSocket /ws/plant funcional, snapshot inicial + eventos. Tiles con color por estado, tickers, KpiOverview y dashboard completo."],
            ["O4. Modelo ML con Random Forest y XGBoost", "Completo (sobre datos sintéticos)", "ml/train.py + ml/features.py (642 LOC combinadas). Modelo serializado, manifest JSON, 4 endpoints REST. Pipeline CRISP-DM ejecutado punta a punta."],
            ["O5. Asistente conversacional", "Parcial", "8 tools + schemas implementados, modo fallback completo y testeado (31 tests). Modo LLM implementado pero no validado: requiere API key real de Anthropic, pendiente para E3."],
            ["O6. Contenerización con Docker Compose", "Completo", "docker-compose.yml de 142 LOC orquesta 5 servicios con healthchecks, volúmenes nombrados, red interna, restart policy."],
            ["O7. Evaluación con usuarios reales", "Pendiente (E3)", "Sin avance esperado en E2. Se programa para semanas 11–14: aplicación de SUS, encuesta satisfacción, registro pre-post."],
        ],
    )

    add_heading(doc, "11.2. Métricas de implementación", level=2)
    add_table(
        doc,
        ["Métrica", "Valor"],
        [
            ["Líneas de código backend (Python)", "~5,400 LOC en app/ + scripts/ + ml/"],
            ["Líneas de código frontend (JSX/JS)", "~3,800 LOC en src/"],
            ["Endpoints REST documentados (Swagger)", "46"],
            ["Canales WebSocket", "1 (/ws/plant)"],
            ["Modelos SQLAlchemy", "14 (más enums)"],
            ["Servicios de negocio", "7"],
            ["Vistas (páginas) frontend", "8"],
            ["Componentes React", "40+"],
            ["Tests automatizados", "197 (167 backend + 30 frontend)"],
            ["Tests pasando al cierre de E2", "197 / 197 (100%)"],
            ["Tasa de cobertura de archivos con docstring", "100%"],
            ["Tablas con migraciones Alembic", "14 (revisión 22d436dd1c32)"],
            ["Modelo ML serializado y cargado en runtime", "Sí (delay-xgb-20260501232338)"],
            ["Healthchecks Docker activos", "Postgres, Redis, Backend (3/5)"],
        ],
    )

    add_heading(doc, "11.3. Pendientes prioritarios para Entrega 3", level=2)
    add_numbered(
        doc,
        [
            "Configurar la API key de Anthropic en producción y validar el modo LLM del chatbot con consultas reales de operación (al menos 50 preguntas predefinidas).",
            "Generar un JWT_SECRET_KEY de producción (64 caracteres aleatorios criptográficamente seguros) y rotar el placeholder actual. Documentar el procedimiento en el manual técnico.",
            "Reemplazar el seed sintético por una carga ETL de datos reales del ERP de al menos 6 meses, reentrenar el modelo y validar la hipótesis H2 (F1 ≥ 0.80).",
            "Diseñar y aplicar el cuestionario SUS y la encuesta de satisfacción con al menos 5 operarios y 2 supervisores reales (validar H4: SUS ≥ 70).",
            "Cronometrar 10 consultas predefinidas pre-implementación y post-implementación con los mismos usuarios para validar H1 (reducción >95%).",
            "Generar el conjunto de pruebas de 50 preguntas para el chatbot y medir la tasa de respuestas correctas por 3 evaluadores (validar H3: ≥85%).",
            "Configurar HTTPS terminado en Nginx con certificados Let's Encrypt o auto-firmados para validar RNF-04 en piloto.",
            "Capturar las imágenes del Anexo D (mockups y vistas reales) e insertarlas en el documento final.",
            "Producir manuales técnico, de despliegue y de usuario.",
            "Hacer el commit inicial del repositorio y publicar la rama main en GitHub para el jurado (actualmente el repo no tiene commits aún).",
        ],
    )
    add_pagebreak(doc)


# -----------------------------------------------------------------------------
# 12. RIESGOS Y MITIGACIONES
# -----------------------------------------------------------------------------
def write_riesgos(doc):
    add_heading(doc, "12. RIESGOS Y MITIGACIONES", level=1)
    add_para(
        doc,
        "La identificación temprana de riesgos es uno de los entregables más "
        "valorados por el jurado en una entrega intermedia. Esta sección recopila "
        "trece riesgos identificados durante la fase de diseño e implementación, "
        "clasificados por probabilidad e impacto, con su estrategia de mitigación. "
        "La probabilidad e impacto se valoran en escala baja/media/alta."
    )
    add_table(
        doc,
        ["ID", "Riesgo", "Prob.", "Impacto", "Mitigación"],
        [
            ["R-01", "Ausencia de datos reales del ERP para entrenar el modelo y bajos F1 actuales (0.21).", "Alta", "Alto", "Datos sintéticos generados por seed para validar pipeline. En E3, coordinar exportación de 6 meses reales del ERP. Plan B: documentar el bajo F1 como limitación y reentrenar en piloto post-grado."],
            ["R-02", "Costos o cortes de la API de Claude (Anthropic) durante demos.", "Media", "Medio", "Modo fallback heurístico cubre los 8 escenarios principales sin LLM. Tests automatizados garantizan su funcionamiento. Plan B alternativo: usar GPT-4 o LLaMA local."],
            ["R-03", "Indisponibilidad de operarios reales para pruebas SUS en planta.", "Media", "Alto", "Identificar al menos 5 operarios y 2 supervisores con anticipación, coordinar con jefe de planta. Plan B: pruebas remotas vía videollamada con operarios."],
            ["R-04", "Cambios de requisitos durante la fase de validación (E3).", "Baja", "Medio", "Documentar requisitos firmados al cierre de E2. Cambios menores se aceptan como mejoras opcionales; cambios mayores se documentan como trabajo futuro."],
            ["R-05", "Pérdida de código ante un fallo de hardware: el repositorio aún no tiene commits.", "Media", "Crítico", "Hacer el commit inicial inmediatamente y push a GitHub privado. Activar backups automáticos del workspace local."],
            ["R-06", "Vulnerabilidades del JWT por uso de placeholder.", "Alta", "Alto", "Rotar JWT_SECRET_KEY a 64 caracteres aleatorios antes de cualquier despliegue real. Documentar en .env.example el comando para generarlo."],
            ["R-07", "Inserción accidental de IMP-03 u otras máquinas durante demos rompiendo tests.", "Media", "Bajo", "Tests refactorizados para tolerar máquinas extra (validan que las del seed estén presentes, pero aceptan adicionales). Resuelto en bloque actual."],
            ["R-08", "Drift del modelo ML al recibir datos del ERP con distribución diferente al sintético.", "Alta", "Medio", "Documentar el procedimiento de reentrenamiento. Monitor offline de AUC-ROC con ventana móvil. Reentrenar al detectar drift > 5 puntos."],
            ["R-09", "Latencia inesperada del WebSocket en redes corporativas con proxies estrictos.", "Media", "Medio", "Validar wss:// en piloto. Plan B: caída a polling cada 5 segundos si la WS pierde conexión durante más de 30 s."],
            ["R-10", "Sobrecosto en horas por sub-estimación de la fase de validación.", "Media", "Alto", "Bloque temporal explícito (semanas 11–14) en el cronograma. Plan B: reducir muestra a mínimo aceptable (3 operarios, 1 supervisor)."],
            ["R-11", "Migración Alembic conflictiva al reactivar el ETL real.", "Baja", "Medio", "Política de migraciones lineales (sin merges). Backup de la BD antes de cualquier `alembic upgrade head` en producción."],
            ["R-12", "Hallazgo durante la sustentación: ausencia de manuales o documentación deficiente.", "Media", "Alto", "Producir tres manuales (técnico, despliegue, usuario) entre semanas 14–16. Revisión cruzada con asesor."],
            ["R-13", "Falla de hardware del PC del estudiante durante la sustentación.", "Baja", "Crítico", "Imagen Docker exportada en USB. Repositorio público disponible. Despliegue alternativo en VPS con un comando."],
        ],
    )
    add_pagebreak(doc)


# -----------------------------------------------------------------------------
# 13. BIBLIOGRAFÍA
# -----------------------------------------------------------------------------
def write_biblio(doc):
    add_heading(doc, "13. BIBLIOGRAFÍA", level=1)
    add_para(
        doc,
        "Las referencias siguientes amplían el cuerpo bibliográfico de la Entrega 1 con "
        "fuentes técnicas específicas para las decisiones de diseño documentadas en "
        "esta entrega. Las referencias del marco teórico permanecen vigentes y se "
        "remiten a la sección 9 de la Entrega 1.",
    )
    refs = [
        "Anaya, V., Alberti, E., & Scivoletto, G. (2024). A Manufacturing Digital Twin Framework. En J. Soldatos (Ed.), Artificial Intelligence in Manufacturing. Springer. https://doi.org/10.1007/978-3-031-46452-2_10",
        "Bangor, A., Kortum, P. T., & Miller, J. T. (2009). Determining what individual SUS scores mean: Adding an adjective rating scale. Journal of Usability Studies, 4(3), 114–123.",
        "Brooke, J. (1996). SUS: A 'quick and dirty' usability scale. En P. W. Jordan, B. Thomas, B. A. Weerdmeester & A. L. McClelland (Eds.), Usability Evaluation in Industry. Taylor & Francis.",
        "Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, 785–794. https://doi.org/10.1145/2939672.2939785",
        "Dai, D., Zhao, B., Yu, Z., Franciosa, P., & Ceglarek, D. (2025). Generative and Predictive AI for digital twin systems in manufacturing. Frontiers in Artificial Intelligence, 8, 1655470.",
        "Digital Twins integrated with Artificial Intelligence: A review across hierarchical manufacturing system levels. (2026). Sensors, 26(1), 124. https://doi.org/10.3390/s26010124",
        "FastAPI Documentation. (2026). FastAPI Framework, high performance, easy to learn, fast to code, ready for production. https://fastapi.tiangolo.com/",
        "Galdino, S., et al. (2025). Large Language Model-Based Cognitive Assistants for Quality Management Systems in Manufacturing: A Requirement Analysis. Engineering Reports.",
        "GitHub. (2026). GitHub Actions Documentation. https://docs.github.com/en/actions",
        "IBM. (2025). 4 ways generative AI addresses manufacturing challenges. IBM Think. https://www.ibm.com/think/topics/generative-ai-for-manufacturing",
        "Kritzinger, W., Karner, M., Traar, G., Henjes, J., & Sihn, W. (2018). Digital Twin in manufacturing: A categorical literature review. IFAC-PapersOnLine, 51(11), 1016–1022.",
        "LangChain Documentation. (2026). LangChain: Building applications with LLMs through composability. https://python.langchain.com/",
        "Machine learning approach for predicting production delays: a quarry company case study. (2022). Journal of Big Data, 9, 91. https://doi.org/10.1186/s40537-022-00644-w",
        "Pedregosa, F., et al. (2011). Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research, 12, 2825–2830.",
        "PostgreSQL Global Development Group. (2026). PostgreSQL 16 Documentation. https://www.postgresql.org/docs/16/",
        "Predictive maintenance in industrial systems: an XGBoost-based approach. (2025). Journal of Industrial and Production Engineering, 42(8), 876–899.",
        "React Team. (2026). React 18 Documentation. https://react.dev/",
        "Sauro, J., & Lewis, J. R. (2016). Quantifying the User Experience: Practical Statistics for User Research (2nd ed.). Morgan Kaufmann.",
        "Shao, G. (2024). Manufacturing Digital Twin Standards. NIST.",
        "Tailwind Labs. (2026). Tailwind CSS Documentation. https://tailwindcss.com/docs",
        "Villegas, M., et al. (2025). Digital twins in manufacturing: A unified conceptual framework. Robotics and Computer-Integrated Manufacturing.",
        "Vite. (2026). Vite — Next Generation Frontend Tooling. https://vitejs.dev/",
        "World Economic Forum. (2024). Why Large Language Models (LLMs) are the future of manufacturing.",
    ]
    for r in refs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(-0.75)
        p.paragraph_format.left_indent = Cm(0.75)
        p.add_run(r).font.size = Pt(10)
    add_pagebreak(doc)


# -----------------------------------------------------------------------------
# ANEXOS
# -----------------------------------------------------------------------------
def write_anexos(doc):
    add_heading(doc, "ANEXOS", level=1)

    # ANEXO A
    add_heading(doc, "ANEXO A. Diccionario de datos completo", level=2)
    add_para(
        doc,
        "Este anexo consolida la totalidad de campos por tabla, incluyendo aquellos "
        "que no se detallaron en la sección 5.3 por brevedad. Cada fila incluye tipo "
        "de dato, restricciones e interpretación funcional."
    )
    # users
    add_para(doc, "users", bold=True)
    add_table(
        doc,
        ["Campo", "Tipo", "Restricciones", "Descripción"],
        [
            ["id", "INTEGER", "PK", "—"],
            ["username", "VARCHAR(64)", "UNIQUE, NOT NULL, indexed", "Login del usuario."],
            ["full_name", "VARCHAR(128)", "NULL", "Nombre completo legible."],
            ["password_hash", "VARCHAR(255)", "NOT NULL", "Hash bcrypt de la contraseña."],
            ["role", "ENUM(UserRole)", "NOT NULL", "admin | supervisor | operator."],
            ["machine_id", "INTEGER", "FK → machines.id, NULL", "Máquina asignada (sólo operarios)."],
            ["is_active", "BOOLEAN", "NOT NULL, default TRUE", "Para soft-delete."],
            ["created_at", "TIMESTAMPTZ", "NOT NULL, default now()", "—"],
        ],
    )
    # shifts
    add_para(doc, "shifts", bold=True)
    add_table(
        doc,
        ["Campo", "Tipo", "Restricciones", "Descripción"],
        [
            ["id", "INTEGER", "PK", "—"],
            ["name", "ENUM(ShiftName)", "NOT NULL", "turno_1 | turno_2 | turno_3."],
            ["start_time", "TIME", "NOT NULL", "Hora de inicio del turno."],
            ["end_time", "TIME", "NOT NULL", "Hora de fin (sigue al inicio modulo 24h)."],
        ],
    )
    # order_operations
    add_para(doc, "order_operations", bold=True)
    add_table(
        doc,
        ["Campo", "Tipo", "Restricciones", "Descripción"],
        [
            ["id", "INTEGER", "PK", "—"],
            ["order_id", "INTEGER", "FK → production_orders.id, NOT NULL", "—"],
            ["machine_id", "INTEGER", "FK → machines.id, NOT NULL", "—"],
            ["sequence", "INTEGER", "NOT NULL, indexed", "Posición en la ruta (1=Tubulado, 2=Impresión, etc.)."],
            ["status", "ENUM(OperationStatus)", "NOT NULL", "pending|ready|in_progress|completed|blocked."],
            ["quantity_in", "INTEGER", "NOT NULL, default 0", "Unidades que entraron a la operación."],
            ["quantity_out", "INTEGER", "NOT NULL, default 0", "Unidades que salieron (para yield)."],
            ["scrap_kg", "FLOAT", "NOT NULL, default 0.0", "Scrap generado en kg."],
            ["scrap_reason", "ENUM(ScrapReason)", "NULL", "quality_defect|setup_loss|material_break|other."],
            ["planned_start", "TIMESTAMPTZ", "NOT NULL", "—"],
            ["planned_end", "TIMESTAMPTZ", "NOT NULL", "—"],
            ["actual_start", "TIMESTAMPTZ", "NULL", "Fecha/hora real de inicio."],
            ["actual_end", "TIMESTAMPTZ", "NULL", "Fecha/hora real de fin."],
            ["operator_id", "INTEGER", "FK → users.id, NULL", "Operario que ejecuta."],
            ["shift", "ENUM(ShiftName)", "NULL", "Turno de ejecución."],
            ["created_at", "TIMESTAMPTZ", "NOT NULL, default now()", "—"],
        ],
    )
    # production_events
    add_para(doc, "production_events", bold=True)
    add_table(
        doc,
        ["Campo", "Tipo", "Restricciones", "Descripción"],
        [
            ["id", "INTEGER", "PK", "—"],
            ["machine_id", "INTEGER", "FK → machines.id, NOT NULL", "—"],
            ["order_id", "INTEGER", "FK → production_orders.id, NULL", "—"],
            ["operation_id", "INTEGER", "FK → order_operations.id, NULL", "—"],
            ["user_id", "INTEGER", "FK → users.id, NULL", "Operario que registra."],
            ["event_type", "ENUM(EventType)", "NOT NULL", "start_operation|stop|format_change|incident|partial_report|end_operation."],
            ["description", "VARCHAR(500)", "NULL", "Descripción libre."],
            ["quantity", "INTEGER", "NULL", "Cantidad reportada (en partial_report o end_operation)."],
            ["scrap_kg", "FLOAT", "NULL", "Kg de scrap del evento."],
            ["scrap_reason", "ENUM(ScrapReason)", "NULL", "—"],
            ["timestamp", "TIMESTAMPTZ", "NOT NULL, default now()", "—"],
        ],
    )
    # oee_records
    add_para(doc, "oee_records", bold=True)
    add_table(
        doc,
        ["Campo", "Tipo", "Restricciones", "Descripción"],
        [
            ["id", "INTEGER", "PK", "—"],
            ["machine_id", "INTEGER", "FK → machines.id, NOT NULL", "—"],
            ["shift_id", "INTEGER", "FK → shifts.id, NOT NULL", "—"],
            ["date", "DATE", "NOT NULL, indexed", "Fecha del registro."],
            ["availability", "FLOAT", "NOT NULL", "0.0–1.0."],
            ["performance", "FLOAT", "NOT NULL", "0.0–1.0."],
            ["quality", "FLOAT", "NOT NULL", "0.0–1.0."],
            ["oee_value", "FLOAT", "NOT NULL", "= A × P × Q."],
        ],
    )
    # ml_predictions
    add_para(doc, "ml_predictions", bold=True)
    add_table(
        doc,
        ["Campo", "Tipo", "Restricciones", "Descripción"],
        [
            ["id", "INTEGER", "PK", "—"],
            ["order_id", "INTEGER", "FK → production_orders.id, NOT NULL", "—"],
            ["delay_probability", "FLOAT", "NOT NULL", "0.0–1.0."],
            ["predicted_delay_hours", "FLOAT", "NOT NULL", "Estimación en horas."],
            ["features_json", "JSON", "NULL", "Snapshot de features usados (para auditoría)."],
            ["model_version", "VARCHAR(32)", "NOT NULL", "Ej. 'delay-xgb-20260501232338'."],
            ["created_at", "TIMESTAMPTZ", "NOT NULL", "—"],
        ],
    )
    # etl_loads
    add_para(doc, "etl_loads", bold=True)
    add_table(
        doc,
        ["Campo", "Tipo", "Restricciones", "Descripción"],
        [
            ["id", "INTEGER", "PK", "—"],
            ["filename", "VARCHAR(255)", "NOT NULL", "—"],
            ["kind", "ENUM(ETLLoadKind)", "NOT NULL", "production_orders|confirmations|materials|shipments."],
            ["status", "ENUM(ETLLoadStatus)", "NOT NULL", "running|success|partial|failed."],
            ["uploaded_by_id", "INTEGER", "FK → users.id, NULL", "—"],
            ["uploaded_at", "TIMESTAMPTZ", "NOT NULL", "—"],
            ["rows_total", "INTEGER", "NOT NULL, default 0", "—"],
            ["rows_inserted", "INTEGER", "NOT NULL, default 0", "—"],
            ["rows_updated", "INTEGER", "NOT NULL, default 0", "—"],
            ["rows_skipped", "INTEGER", "NOT NULL, default 0", "—"],
            ["rows_failed", "INTEGER", "NOT NULL, default 0", "—"],
            ["duration_ms", "INTEGER", "NOT NULL, default 0", "—"],
            ["error_log", "JSON", "NULL", "Errores por fila."],
        ],
    )
    add_pagebreak(doc)

    # ANEXO B - Casos de uso detallados
    add_heading(doc, "ANEXO B. Fichas detalladas de casos de uso", level=2)

    # B.1 CU-01
    add_heading(doc, "B.1. CU-01 — Autenticarse en el sistema", level=3)
    add_table(
        doc,
        ["Atributo", "Valor"],
        [
            ["Actor primario", "Cualquier usuario (operario, supervisor, administrador)"],
            ["Precondición", "El usuario tiene credenciales activas en la base de datos."],
            ["Disparador", "El usuario abre la URL raíz y es redirigido a /login."],
            ["Flujo básico",
             "1. El usuario ingresa username y password. "
             "2. Pulsa el botón 'Iniciar sesión'. "
             "3. El frontend invoca POST /api/auth/login con form-data. "
             "4. El backend valida credenciales con bcrypt. "
             "5. El backend emite JWT firmado con HS256. "
             "6. El frontend guarda el token en localStorage. "
             "7. El frontend redirige según el rol: operario→/operator, supervisor→/dashboard, admin→/admin."],
            ["Flujo alternativo", "F4a. Credenciales inválidas → 401, mensaje 'Credenciales inválidas'. "
             "F4b. Usuario inactivo → 403, mensaje 'Cuenta inactiva'."],
            ["Postcondición", "El navegador del usuario tiene un JWT válido en localStorage."],
            ["Cumple RF", "RF-01, RF-02"],
        ],
    )

    # B.2 CU-04
    add_heading(doc, "B.2. CU-04 — Iniciar operación", level=3)
    add_table(
        doc,
        ["Atributo", "Valor"],
        [
            ["Actor primario", "Operario"],
            ["Precondición", "Existe una operación con status READY asignada a la máquina del operario."],
            ["Disparador", "El operario pulsa 'Iniciar' en MachineActionButtons sobre la operación de la cola."],
            ["Flujo básico",
             "1. La UI muestra un Modal pidiendo confirmación. "
             "2. El operario confirma. "
             "3. El frontend invoca POST /api/operations/{operation_id}/start. "
             "4. El backend valida que la operación esté READY y que el operario coincida con la asignación. "
             "5. El backend marca operation.status=IN_PROGRESS, actual_start=now(). "
             "6. El backend persiste un ProductionEvent de tipo start_operation. "
             "7. El backend emite por WebSocket al canal /ws/plant un mensaje state_change. "
             "8. La UI del operario y la UI del supervisor se actualizan."],
            ["Flujo alternativo", "F4a. Operación no es READY → 409 Conflict. "
             "F4b. Operación no asignada al operario → 403 Forbidden."],
            ["Postcondición", "Operation.status=IN_PROGRESS, machine.status=RUNNING."],
            ["Cumple RF", "RF-03, RF-04"],
        ],
    )

    # B.3 CU-05
    add_heading(doc, "B.3. CU-05 — Registrar evento de máquina", level=3)
    add_table(
        doc,
        ["Atributo", "Valor"],
        [
            ["Actor primario", "Operario"],
            ["Precondición", "El operario está autenticado."],
            ["Disparador", "El operario pulsa uno de los botones rápidos: Parada, Cambio de formato, Incidencia."],
            ["Flujo básico",
             "1. La UI abre un Modal con el formulario adecuado al tipo de evento. "
             "2. El operario describe brevemente y, si aplica, indica scrap_kg y scrap_reason. "
             "3. El frontend invoca POST /api/events con event_type, description, machine_id, order_id, operation_id. "
             "4. El backend valida, persiste en production_events y, si es un STOP, actualiza machine.status=DOWN. "
             "5. El backend emite el evento por /ws/plant. "
             "6. La UI muestra un toast 'Evento registrado' y refresca el RecentEventsLog."],
            ["Flujo alternativo", "F3a. Validación falla → 422 Unprocessable Entity con detalle por campo en español."],
            ["Postcondición", "Existe un nuevo registro en production_events vinculado a la máquina y al operario."],
            ["Cumple RF", "RF-04, RF-12"],
        ],
    )

    # B.4 CU-06
    add_heading(doc, "B.4. CU-06 — Reportar avance parcial", level=3)
    add_table(
        doc,
        ["Atributo", "Valor"],
        [
            ["Actor primario", "Operario"],
            ["Precondición", "Existe una operación IN_PROGRESS asignada al operario."],
            ["Disparador", "El operario pulsa 'Reportar avance' en MachineActionButtons."],
            ["Flujo básico",
             "1. La UI abre un Modal con el formulario: cantidad producida (parcial). "
             "2. El operario ingresa la cantidad. "
             "3. El frontend invoca POST /api/operations/{operation_id}/report con la cantidad. "
             "4. El backend acumula en operation.quantity_out y refleja la suma en order.quantity_produced. "
             "5. El backend persiste un ProductionEvent de tipo partial_report. "
             "6. El backend emite por WS un mensaje progress_update."],
            ["Flujo alternativo", "F4a. La cantidad reportada excede quantity_in → 422."],
            ["Postcondición", "Order.quantity_produced reflejado en la UI del supervisor en menos de 1 segundo."],
            ["Cumple RF", "RF-04, RF-13"],
        ],
    )

    # B.5 CU-08
    add_heading(doc, "B.5. CU-08 — Visualizar Digital Twin", level=3)
    add_table(
        doc,
        ["Atributo", "Valor"],
        [
            ["Actor primario", "Supervisor / Administrador"],
            ["Precondición", "Usuario autenticado con rol supervisor o admin."],
            ["Disparador", "El usuario navega a /supervisor."],
            ["Flujo básico",
             "1. La SPA renderiza PlantMap. "
             "2. PlantMap inicia conexión WebSocket a /ws/plant?token=<JWT>. "
             "3. El backend valida el token y emite un mensaje snapshot con todas las máquinas. "
             "4. PlantMap renderiza un MachineTile por cada máquina. "
             "5. Cuando llega un mensaje state_change o event, el tile correspondiente se actualiza."],
            ["Flujo alternativo", "F2a. Token inválido → WS cerrado con código 1008."],
            ["Postcondición", "El supervisor ve la planta en tiempo real."],
            ["Cumple RF", "RF-06"],
        ],
    )

    # B.6 CU-09
    add_heading(doc, "B.6. CU-09 — Consultar dashboard de KPIs", level=3)
    add_table(
        doc,
        ["Atributo", "Valor"],
        [
            ["Actor primario", "Supervisor / Administrador"],
            ["Precondición", "Usuario autenticado con rol supervisor o admin."],
            ["Disparador", "El usuario navega a /dashboard."],
            ["Flujo básico",
             "1. La SPA invoca GET /api/dashboard/overview, /oee-trend, /machine-ranking, /alerts, /scrap-by-machine, etc. (siete endpoints en paralelo). "
             "2. El backend computa cada KPI en su servicio. "
             "3. La SPA renderiza nueve widgets con Recharts. "
             "4. El usuario puede ajustar el rango de días en la URL (?days=N) para refiltrar."],
            ["Flujo alternativo", "F2a. Si days es < 1 o > 365 → 422."],
            ["Postcondición", "El supervisor visualiza los KPIs de la planta."],
            ["Cumple RF", "RF-07, RF-08"],
        ],
    )

    # B.7 CU-11
    add_heading(doc, "B.7. CU-11 — Conversar con el asistente IA", level=3)
    add_table(
        doc,
        ["Atributo", "Valor"],
        [
            ["Actor primario", "Supervisor / Administrador"],
            ["Precondición", "Usuario autenticado."],
            ["Disparador", "El usuario navega a /chat y escribe una pregunta."],
            ["Flujo básico",
             "1. El usuario escribe '¿Cuál es el OEE de la planta?' y pulsa Enter. "
             "2. La SPA invoca POST /api/chat/message con message + history. "
             "3. El backend evalúa si hay API key válida (is_llm_available()). "
             "4. Si hay key → invoca ChatAnthropic con bind_tools. El modelo decide invocar get_oee_data, recibe el resultado, redacta la respuesta. "
             "5. Si no hay key → router heurístico mapea 'OEE' a get_oee_data y formatea respuesta. "
             "6. La SPA muestra la respuesta + tool_calls + badge del modo."],
            ["Flujo alternativo", "F4a. LLM falla → fallback con campo error documentado. "
             "F2a. message > 2000 caracteres → 422."],
            ["Postcondición", "El usuario recibe la respuesta en español con datos reales de la BD."],
            ["Cumple RF", "RF-09"],
        ],
    )
    add_pagebreak(doc)

    # ANEXO C - Inventario de endpoints REST
    add_heading(doc, "ANEXO C. Inventario completo de endpoints REST", level=2)
    add_para(
        doc,
        "El siguiente inventario corresponde al estado actual de la API descrita por "
        "la especificación OpenAPI generada por FastAPI. El total es de 46 endpoints "
        "REST más 1 canal WebSocket (/ws/plant)."
    )
    add_table(
        doc,
        ["Método", "Ruta", "Rol mínimo", "Propósito"],
        [
            ["GET", "/", "Público", "Bienvenida del backend."],
            ["GET", "/health", "Público", "Health check rápido."],
            ["GET", "/api/health", "Público", "Health check API + timestamp."],
            ["POST", "/api/auth/login", "Público", "Emite JWT (form-data OAuth2)."],
            ["GET", "/api/auth/me", "Auth", "Datos del usuario actual."],
            ["GET", "/api/machines", "Auth", "Listar máquinas (filtro por type opcional)."],
            ["POST", "/api/machines", "Admin", "Crear máquina."],
            ["GET", "/api/machines/{id}", "Auth", "Detalle de máquina."],
            ["PATCH", "/api/machines/{id}", "Supervisor+", "Actualizar (status, location)."],
            ["DELETE", "/api/machines/{id}", "Admin", "Eliminar máquina."],
            ["GET", "/api/orders", "Auth", "Listar órdenes (filtros: status, machine_id, etc.)."],
            ["POST", "/api/orders", "Supervisor+", "Crear orden."],
            ["GET", "/api/orders/{id}", "Auth", "Detalle de orden."],
            ["PATCH", "/api/orders/{id}", "Supervisor+", "Actualizar orden."],
            ["DELETE", "/api/orders/{id}", "Admin", "Eliminar."],
            ["GET", "/api/orders/{id}/operations", "Auth", "Listar operaciones de la orden."],
            ["GET", "/api/operations", "Auth", "Listar operaciones (filtros)."],
            ["GET", "/api/operations/{id}", "Auth", "Detalle de operación."],
            ["POST", "/api/operations/{id}/start", "Operario+", "Iniciar operación."],
            ["POST", "/api/operations/{id}/report", "Operario+", "Reportar avance."],
            ["POST", "/api/operations/{id}/complete", "Operario+", "Cerrar operación."],
            ["GET", "/api/events", "Auth", "Listar eventos."],
            ["POST", "/api/events", "Operario+", "Crear evento."],
            ["GET", "/api/events/{id}", "Auth", "Detalle de evento."],
            ["GET", "/api/dashboard/overview", "Supervisor+", "Resumen plant_oee, A/P/Q, conteos."],
            ["GET", "/api/dashboard/oee-trend", "Supervisor+", "Tendencia 30 días."],
            ["GET", "/api/dashboard/production-by-shift", "Supervisor+", "Producción por turno."],
            ["GET", "/api/dashboard/order-fulfillment", "Supervisor+", "Cumplimiento de órdenes."],
            ["GET", "/api/dashboard/machine-ranking", "Supervisor+", "Ranking por OEE."],
            ["GET", "/api/dashboard/alerts", "Supervisor+", "Alertas predictivas."],
            ["GET", "/api/dashboard/scrap-by-machine", "Supervisor+", "Pareto de scrap."],
            ["GET", "/api/dashboard/yield-by-operation", "Supervisor+", "Yield por máquina."],
            ["GET", "/api/dashboard/wip", "Supervisor+", "Work In Progress."],
            ["GET", "/api/predictions/model-info", "Supervisor+", "Versión del modelo + métricas."],
            ["GET", "/api/predictions/feature-importance", "Supervisor+", "Top features del modelo."],
            ["POST", "/api/predictions/predict/{order_id}", "Supervisor+", "Predecir una orden."],
            ["POST", "/api/predictions/predict-active", "Supervisor+", "Predecir todas las activas."],
            ["POST", "/api/predictions/reload", "Admin", "Recargar modelo desde disco."],
            ["GET", "/api/chat/status", "Auth", "Estado del chatbot."],
            ["POST", "/api/chat/message", "Auth", "Conversar."],
            ["POST", "/api/etl/upload", "Admin", "Subir CSV."],
            ["GET", "/api/etl/status", "Admin", "Bitácora de cargas."],
            ["GET", "/api/etl/status/{id}", "Admin", "Detalle de una carga."],
            ["GET", "/api/etl/sample-csv/{kind}", "Admin", "Descargar plantilla."],
            ["GET", "/api/users", "Admin", "Listar usuarios."],
            ["POST", "/api/users", "Admin", "Crear usuario."],
            ["GET", "/api/users/{id}", "Admin", "Detalle."],
            ["PATCH", "/api/users/{id}", "Admin", "Editar."],
            ["DELETE", "/api/users/{id}", "Admin", "Desactivar."],
            ["POST", "/api/users/{id}/assign-machine", "Admin", "Asignar máquina."],
            ["POST", "/api/users/{id}/reset-password", "Admin", "Resetear password."],
            ["GET", "/api/admin/audit", "Admin", "Auditoría."],
            ["GET", "/api/admin/system/health", "Admin", "Salud del sistema."],
            ["GET", "/api/admin/system/ml-status", "Admin", "Estado del modelo."],
            ["GET", "/api/admin/system/settings", "Admin", "Listar settings."],
            ["PATCH", "/api/admin/system/settings/{key}", "Admin", "Editar setting."],
            ["WS", "/ws/plant", "Supervisor+", "Canal WebSocket del Digital Twin."],
        ],
    )
    add_pagebreak(doc)

    # ANEXO D
    add_heading(doc, "ANEXO D. Mockups e imágenes (placeholder)", level=2)
    add_para(
        doc,
        "Este anexo está reservado para las capturas de pantalla del sistema "
        "operativo. Se recomienda incluir doce imágenes que ilustren el recorrido "
        "completo por las ocho vistas: Login, Dashboard (con datos llenos), "
        "OperatorView (con operación en curso), SupervisorView (Digital Twin con "
        "todas las máquinas), OrderTraceView, Chat (con varios turnos de "
        "conversación), ETL (antes/después de cargar archivo) y Admin (cada "
        "pestaña). Las capturas se obtienen levantando el sistema con "
        "`docker compose up -d`, ingresando con cada uno de los tres roles, y "
        "usando la herramienta de captura del sistema operativo. Para esta entrega, "
        "el lector puede acceder al sistema en vivo en http://localhost tras "
        "ejecutar el seed y validar visualmente cada flujo descrito en la sección 6."
    )

    # ANEXO E - GitHub Actions skeleton
    add_heading(doc, "ANEXO E. Plantilla de CI con GitHub Actions", level=2)
    add_para(
        doc,
        "El siguiente esqueleto de workflow .github/workflows/ci.yml se recomienda "
        "para la Entrega 3. Levanta los servicios necesarios, instala "
        "dependencias y ejecuta las suites de pruebas en cada push y pull request."
    )
    snippet = """name: CI
on:
  push:
    branches: [main]
  pull_request:
    branches: [main]
jobs:
  backend-tests:
    runs-on: ubuntu-latest
    services:
      postgres:
        image: postgres:16-alpine
        env:
          POSTGRES_USER: smartsack
          POSTGRES_PASSWORD: smartsack_dev_password
          POSTGRES_DB: smartsack
        ports: ['5432:5432']
        options: >-
          --health-cmd pg_isready --health-interval 5s
          --health-timeout 3s --health-retries 5
      redis:
        image: redis:7-alpine
        ports: ['6379:6379']
    steps:
      - uses: actions/checkout@v4
      - uses: actions/setup-python@v5
        with: { python-version: '3.11' }
      - run: pip install -r backend/requirements.txt
      - run: cd backend && alembic upgrade head
      - run: cd backend && python -m scripts.seed
      - run: cd backend && pytest --tb=short -q
  frontend-tests:
    runs-on: ubuntu-latest
    steps:
      - uses: actions/checkout@v4
      - uses: actions/setup-node@v4
        with: { node-version: '20' }
      - run: cd frontend && npm ci
      - run: cd frontend && npm run test -- --run
"""
    p = doc.add_paragraph()
    p.add_run(snippet).font.name = "Courier New"
    for run in p.runs:
        run.font.size = Pt(8)
    add_pagebreak(doc)


# -----------------------------------------------------------------------------
# MAIN
# -----------------------------------------------------------------------------
def main() -> Path:
    doc = Document()
    # Configurar fuentes por defecto
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    write_cover(doc)
    write_resumen(doc)
    write_intro(doc)
    write_requisitos(doc)
    write_casos_uso(doc)
    write_arquitectura(doc)
    write_modelo_datos(doc)
    write_interfaz(doc)
    write_ml(doc)
    write_etl(doc)
    write_chatbot(doc)
    write_pruebas(doc)
    write_avance(doc)
    write_riesgos(doc)
    write_biblio(doc)
    write_anexos(doc)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    out = main()
    print(f"OK -> {out} ({out.stat().st_size} bytes)")
